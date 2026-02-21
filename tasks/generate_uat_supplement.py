"""Generate UAT Supplement DOCX covering gaps not in the main UAT docs.

Covers:
- Phase 3 minor gaps (landing page, logout, negative auth, dashboard KPIs)
- Phase 4 - Mobile app flows
- Phase 5 - API direct endpoint validation
- Phase 6 - Security validation
- Phase 7 - Worker & queue validation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from lxml import etree

    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elem.set(qn("w:val"), "clear")
    shading_elem.set(qn("w:color"), "auto")
    shading_elem.set(qn("w:fill"), color_hex)


def make_doc(title, subtitle):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Winnow Job Matching Platform")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  |  February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()
    return doc


def add_scenario(doc, scenario_id, title, benefit, steps, expected, differentiator=None):
    """Add a test scenario with pass/fail checkbox."""
    doc.add_heading(f"{scenario_id}: {title}", level=3)

    p = doc.add_paragraph()
    run = p.add_run("Benefit: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(benefit)
    run.font.size = Pt(9)
    run.font.italic = True

    doc.add_paragraph("Steps:", style="List Bullet")
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Expected Result: ")
    run.bold = True
    run = p.add_run(expected)

    if differentiator:
        p = doc.add_paragraph()
        run = p.add_run("WINNOW DIFFERENTIATOR: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x0A, 0x7E, 0x3F)
        run = p.add_run(differentiator)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x0A, 0x7E, 0x3F)

    # Pass/Fail line
    p = doc.add_paragraph()
    run = p.add_run("Result:  ")
    run.bold = True
    run = p.add_run(
        "[ ] PASS    [ ] FAIL    Tester: ___________    Date: ___________"
    )
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run("Notes: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run("_" * 80)
    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer


def add_checklist_table(doc, scenarios):
    """Add a summary pass/fail checklist table."""
    doc.add_heading("Summary Checklist", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["ID", "Scenario", "Pass/Fail", "Notes"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "1B3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for sid, title in scenarios:
        row = table.add_row()
        row.cells[0].text = sid
        row.cells[1].text = title
        row.cells[2].text = "[ ] P  [ ] F"
        row.cells[3].text = ""
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)


def build_supplement_doc():
    doc = make_doc(
        "UAT Supplement",
        "Coverage Gaps: Web Basics, Mobile, API Direct, Security",
    )

    # --- PURPOSE ---
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This supplement covers test scenarios NOT included in the three main UAT "
        "scripts (Candidate, Employer, Recruiter/Admin). It fills gaps identified "
        "when mapping the PROMPT25 QA checklist (53 steps across 8 phases) against "
        "the existing UAT documents. The main UAT docs cover ~85% of Phase 3 web "
        "flows with excellent depth; this supplement covers the remaining gaps."
    )

    doc.add_heading("Prerequisites", level=1)
    prereqs = [
        "All services running: Docker (Postgres + Redis), API, Worker, Web",
        "Test candidate account: candidate@test.com / TestPass123!",
        "Expo Go app installed on a physical device or emulator for mobile tests",
        "Stripe test mode configured (test key 4242 4242 4242 4242)",
        "Admin token configured (ADMIN_TOKEN env var = dev-admin-token)",
        "At least 10 active jobs and 1 completed onboarding in the database",
    ]
    for p_text in prereqs:
        doc.add_paragraph(p_text, style="List Bullet")

    doc.add_page_break()

    scenarios = []

    # ================================================================
    # SECTION 1: WEB APP MINOR GAPS (Phase 3)
    # ================================================================
    doc.add_heading("Section 1: Web App Flow Gaps", level=1)
    doc.add_paragraph(
        "These scenarios cover Phase 3 items from PROMPT25 that were not included "
        "in the main Candidate UAT doc: landing page verification, standalone "
        "logout, negative authentication cases, and dashboard KPI standalone test."
    )

    sid = "S-01"
    scenarios.append((sid, "Landing Page Load"))
    add_scenario(
        doc,
        sid,
        "Landing Page Load & Visual Verification",
        "First impressions matter -- the landing page must load cleanly with correct branding.",
        [
            "Open http://localhost:3000 in a fresh browser (incognito/private mode)",
            "Verify the landing page loads without errors",
            "Open DevTools (F12) -> Console tab",
            "Confirm NO JavaScript errors or unhandled exceptions in the console",
            "Verify brand colors display correctly (hunter green #1B3A5C, gold accents)",
            "Verify the login form or CTA is visible and accessible",
            "Verify page is responsive: resize browser to mobile width (375px) and confirm layout adapts",
        ],
        "Landing page loads cleanly with correct branding, no console errors, "
        "and responsive layout. Login/signup CTAs are prominent and functional.",
    )

    sid = "S-02"
    scenarios.append((sid, "Standalone Logout"))
    add_scenario(
        doc,
        sid,
        "Logout Flow & Session Cleanup",
        "Users must be able to log out cleanly with session fully invalidated.",
        [
            "Log in with a valid test account",
            "Verify you land on /dashboard",
            "Click the logout button (profile menu, settings, or nav bar)",
            "Verify redirect to the login page",
            "Open DevTools -> Application -> Cookies",
            "Verify the rm_session cookie is cleared or expired",
            "Manually navigate to /dashboard by typing the URL",
            "Verify you are redirected to /login (auth guard works)",
            "Manually navigate to /profile",
            "Verify you are redirected to /login",
            "Manually navigate to /matches",
            "Verify you are redirected to /login",
        ],
        "Logout clears session cookie. All protected routes (/dashboard, /profile, "
        "/matches) redirect to /login when unauthenticated. No cached data visible.",
    )

    sid = "S-03"
    scenarios.append((sid, "Negative Auth Cases"))
    add_scenario(
        doc,
        sid,
        "Login Error Handling & Security",
        "Error messages must be helpful but not leak information about account existence.",
        [
            "Navigate to /login",
            "Enter a VALID email but WRONG password",
            "Click login",
            "Verify an error message appears (e.g., 'Invalid credentials')",
            "Verify the error does NOT say 'wrong password' (information leak)",
            "Clear the form",
            "Enter a NONEXISTENT email (e.g., doesnotexist@fake.dev) with any password",
            "Click login",
            "Verify the SAME generic error message appears (not 'user not found')",
            "Verify no stack traces or technical details are shown",
            "Attempt login with an empty email field -- verify client-side validation prevents submission",
            "Attempt login with an empty password field -- verify client-side validation prevents submission",
        ],
        "Wrong password and nonexistent email produce the same generic error message. "
        "No information leakage about account existence. Client-side validation catches empty fields.",
    )

    sid = "S-04"
    scenarios.append((sid, "Dashboard KPI Standalone"))
    add_scenario(
        doc,
        sid,
        "Dashboard 5 KPI Cards & Pipeline Visualization",
        "The dashboard must display all 5 metric cards with accurate, live data.",
        [
            "Log in as a candidate with completed onboarding and at least some matches",
            "Navigate to /dashboard",
            "Verify 'Dashboard' heading is visible",
            "Verify 5 KPI metric cards are displayed:",
            "  - Profile Completeness (percentage with color: red <50%, amber 50-79%, green 80%+)",
            "  - Qualified Jobs (count of matches scoring 50+)",
            "  - Applications Submitted (count of matches with status 'applied')",
            "  - Interviews Requested (count of matches with status 'interviewing')",
            "  - Offers Received (count of matches with status 'offer')",
            "Click 'Profile Completeness' card -- verify navigation to /profile",
            "Navigate back to /dashboard",
            "Click 'Qualified Jobs' card -- verify navigation to /matches",
            "Navigate back to /dashboard",
            "Verify pipeline/funnel visualization is displayed below the KPI cards",
            "Cross-check: go to /matches, count matches -- verify 'Qualified Jobs' number matches",
            "Change a match status to 'applied', return to /dashboard",
            "Verify 'Applications Submitted' count incremented by 1",
        ],
        "All 5 KPI cards display accurate, live data. Card click navigation works. "
        "Pipeline visualization renders. Metrics update in real time after status changes.",
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 2: MOBILE APP FLOWS (Phase 4)
    # ================================================================
    doc.add_heading("Section 2: Mobile App Flows", level=1)
    doc.add_paragraph(
        "These scenarios cover PROMPT25 Phase 4 (Steps 36-40). The mobile app "
        "is built with Expo/React Native. Test using Expo Go on a physical device "
        "or an iOS/Android emulator."
    )

    p = doc.add_paragraph()
    run = p.add_run("Setup: ")
    run.bold = True
    run = p.add_run(
        "Start the mobile dev server: cd apps/mobile && npx expo start --offline. "
        "Scan the QR code with Expo Go on your device."
    )

    sid = "S-05"
    scenarios.append((sid, "Mobile Login"))
    add_scenario(
        doc,
        sid,
        "Mobile Login & Error Handling",
        "Mobile users must authenticate with the same reliability as web users.",
        [
            "Open the app in Expo Go",
            "Verify the login screen displays with brand styling (hunter green + gold)",
            "Enter valid credentials (candidate@test.com / TestPass123!)",
            "Tap 'Log In'",
            "Verify navigation to the dashboard tab",
            "Log out (via profile tab)",
            "Return to login screen",
            "Enter a wrong password",
            "Tap 'Log In'",
            "Verify an error message appears (not a crash or unhandled exception)",
            "Verify the app remains responsive after the error",
        ],
        "Login succeeds with valid credentials and navigates to dashboard. "
        "Wrong password shows a user-friendly error. App remains stable after auth errors.",
    )

    sid = "S-06"
    scenarios.append((sid, "Mobile Dashboard"))
    add_scenario(
        doc,
        sid,
        "Mobile Dashboard Tab & Metrics",
        "Mobile users see the same key metrics as web users in a mobile-optimized layout.",
        [
            "Log in and land on the dashboard tab",
            "Verify metric cards are displayed (Profile Completeness, Qualified Jobs, etc.)",
            "Verify the plan badge is visible (Free or Pro)",
            "Tap 'View Matches' button (if present)",
            "Verify navigation to the matches tab",
            "Navigate back to dashboard",
            "Verify data matches what the web dashboard shows for the same account",
        ],
        "Dashboard tab shows metric cards with correct data. Plan badge visible. "
        "Navigation to matches works.",
    )

    sid = "S-07"
    scenarios.append((sid, "Mobile Matches List"))
    add_scenario(
        doc,
        sid,
        "Mobile Matches Tab & Pull-to-Refresh",
        "Mobile users can browse and refresh matches with native mobile gestures.",
        [
            "Tap the Matches tab",
            "Verify a list of match cards is displayed",
            "Verify each card shows: job title, company name, score badge, top skills",
            "Pull down on the list to trigger pull-to-refresh",
            "Verify the list reloads (loading indicator appears briefly)",
            "Verify match count is consistent with the web matches page",
            "Scroll through the list to verify pagination or infinite scroll works",
        ],
        "Match cards display with correct data. Pull-to-refresh reloads the list. "
        "Scrolling is smooth and pagination works.",
    )

    sid = "S-08"
    scenarios.append((sid, "Mobile Job Detail"))
    add_scenario(
        doc,
        sid,
        "Mobile Job Detail, Status, & Resume Generation",
        "Mobile users get the full match detail experience optimized for small screens.",
        [
            "From the matches tab, tap a match card",
            "Verify navigation to the job detail screen",
            "Verify match score and interview readiness are displayed prominently",
            "Verify 'Reasons' section shows matched skills with evidence",
            "Verify 'Gaps' section shows missing skills with recommendations",
            "Tap the application status picker",
            "Change status from 'saved' to 'applied'",
            "Verify status change persists (navigate away and return)",
            "Tap 'Generate ATS Resume'",
            "Verify loading indicator appears",
            "Wait for generation to complete",
            "Tap 'Download' or 'Share' button",
            "Verify the native share sheet opens (iOS) or share intent fires (Android)",
        ],
        "Job detail shows full match analysis. Status picker saves changes. "
        "Resume generation completes. Native share sheet works for downloading/sharing.",
    )

    sid = "S-09"
    scenarios.append((sid, "Mobile Profile & Logout"))
    add_scenario(
        doc,
        sid,
        "Mobile Profile Edit, Save, Logout, & Token Persistence",
        "Mobile users can manage preferences and maintain sessions across app restarts.",
        [
            "Tap the Profile tab",
            "Verify editable preference fields are displayed (e.g., remote preference, target roles)",
            "Edit a preference (e.g., change remote preference from 'Remote' to 'Hybrid')",
            "Tap 'Save'",
            "Verify a success confirmation appears",
            "Navigate away from profile tab, then return",
            "Verify the change persisted",
            "Tap 'Log Out'",
            "Verify return to the login screen",
            "Log in again",
            "Verify the edited preference still reflects the saved change",
            "Close the app completely (swipe away from recent apps)",
            "Reopen the app",
            "Verify you are still logged in (token persisted in secure storage)",
        ],
        "Profile edits save correctly. Logout clears the session. Re-login shows "
        "persisted data. App restart maintains authentication via stored token.",
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 3: API DIRECT ENDPOINT VALIDATION (Phase 5)
    # ================================================================
    doc.add_heading("Section 3: API Direct Endpoint Validation", level=1)
    doc.add_paragraph(
        "These scenarios cover PROMPT25 Phase 5 (Steps 41-44). Use Swagger UI "
        "at http://127.0.0.1:8000/docs or curl/PowerShell commands to test "
        "endpoints that aren't easily exercised through the web or mobile UI."
    )

    sid = "S-10"
    scenarios.append((sid, "Auth API Endpoints"))
    add_scenario(
        doc,
        sid,
        "Auth Endpoints: Signup, Login, Me",
        "The auth API must return correct responses and enforce security.",
        [
            "POST /api/auth/signup with JSON body: "
            "{email: 'apitest_supp@winnow.dev', password: 'TestPass123!'}",
            "Verify response: 200 with user_id, email, token fields",
            "POST /api/auth/login with same credentials",
            "Verify response: 200 with user_id, email, token, onboarding_complete fields",
            "Save the token from the login response",
            "GET /api/auth/me with header: Authorization: Bearer <token>",
            "Verify response: 200 with user info (id, email, role)",
            "GET /api/auth/me WITHOUT any auth header",
            "Verify response: 401 Unauthorized",
            "GET /api/auth/me with an invalid/expired token: Bearer invalid-token-123",
            "Verify response: 401 Unauthorized",
        ],
        "Signup and login return tokens. /me returns user info with valid token. "
        "Missing or invalid tokens return 401.",
    )

    sid = "S-11"
    scenarios.append((sid, "Sieve API Endpoints"))
    add_scenario(
        doc,
        sid,
        "Sieve Chat, Triggers, & History Endpoints",
        "The Sieve AI API must return contextual, personalized responses.",
        [
            "Obtain a valid Bearer token by logging in as a candidate with completed onboarding",
            "POST /api/sieve/chat with JSON body: "
            "{message: 'What are my top matches?', conversation_history: []}",
            "  Include header: Authorization: Bearer <token>",
            "Verify response contains: reply text, suggestions array",
            "Verify the reply references actual match data (not generic career advice)",
            "POST /api/sieve/triggers with JSON body: {dismissed_ids: []}",
            "  Include header: Authorization: Bearer <token>",
            "Verify response contains a list of applicable triggers based on user state",
            "GET /api/sieve/history with Authorization header",
            "Verify response contains saved conversation messages",
            "POST /api/sieve/chat WITHOUT auth header",
            "Verify response: 401 Unauthorized",
        ],
        "Chat returns contextual responses referencing actual user data. "
        "Triggers return state-based nudges. History returns saved messages. Auth is enforced.",
    )

    sid = "S-12"
    scenarios.append((sid, "Billing API Endpoint"))
    add_scenario(
        doc,
        sid,
        "Billing Status Endpoint",
        "The billing API must accurately report subscription state and usage.",
        [
            "Obtain a valid Bearer token",
            "GET /api/billing/status with Authorization header",
            "Verify response contains: plan (free/pro), usage counts, limits",
            "Verify usage counts match actual usage "
            "(cross-check with /matches and /tailor counts)",
            "GET /api/billing/status WITHOUT auth header",
            "Verify response: 401 Unauthorized",
        ],
        "Billing status returns accurate plan, usage, and limit information. Auth is enforced.",
    )

    sid = "S-13"
    scenarios.append((sid, "Admin Observability Endpoints"))
    add_scenario(
        doc,
        sid,
        "Admin Observability & Security Posture Endpoints",
        "Admin endpoints must provide full system visibility with proper auth gating.",
        [
            "GET /api/admin/observability/health?admin_token=dev-admin-token",
            "Verify response includes: API status, database status, Redis status, uptime",
            "GET /api/admin/observability/queues?admin_token=dev-admin-token",
            "Verify response includes per-queue stats: queued, started, finished, failed counts",
            "GET /api/admin/security-posture?admin_token=dev-admin-token",
            "Verify response includes checks for: AUTH_SECRET strength, environment mode, "
            "CORS, Stripe key mode",
            "Call any admin endpoint WITHOUT the admin_token parameter",
            "Verify response: 401 or 403 (access denied)",
            "Call any admin endpoint with an INCORRECT admin_token",
            "Verify response: 401 or 403 (access denied)",
        ],
        "Admin endpoints return detailed system information with valid admin token. "
        "Invalid or missing admin tokens are rejected with 401/403.",
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 4: SECURITY VALIDATION (Phase 6)
    # ================================================================
    doc.add_heading("Section 4: Security Validation", level=1)
    doc.add_paragraph(
        "These scenarios cover PROMPT25 Phase 6 (Steps 45-48). Security testing "
        "validates rate limiting, response headers, input validation, and auth "
        "hardening across the platform."
    )

    sid = "S-14"
    scenarios.append((sid, "Rate Limiting"))
    add_scenario(
        doc,
        sid,
        "Rate Limiting on Auth & Chat Endpoints",
        "Rate limiting prevents brute-force attacks and API abuse.",
        [
            "Open a terminal/PowerShell window",
            "Send 11+ rapid POST requests to /api/auth/login with wrong credentials "
            "within 60 seconds:",
            "  Use a loop: for ($i=0; $i -lt 15; $i++) { "
            "curl -X POST http://localhost:8000/api/auth/login ... }",
            "Verify that after the rate limit threshold, the response status is "
            "429 Too Many Requests",
            "Verify the 429 response includes a Retry-After header or descriptive message",
            "Wait 60 seconds for the rate limit window to reset",
            "Obtain a valid Bearer token",
            "Send 31+ rapid POST requests to /api/sieve/chat within 60 seconds",
            "Verify that after the rate limit threshold, the response status is 429",
            "Verify the app remains functional after rate limiting (not crashed)",
        ],
        "Login endpoint returns 429 after 10+ attempts/minute. Sieve chat returns "
        "429 after 30+ messages/minute. Rate limits reset after the window expires. "
        "App remains stable.",
    )

    sid = "S-15"
    scenarios.append((sid, "Security Headers"))
    add_scenario(
        doc,
        sid,
        "Security Response Headers Verification",
        "Proper security headers protect against common web vulnerabilities.",
        [
            "Open the web app at http://localhost:3000",
            "Open DevTools (F12) -> Network tab",
            "Click on any page navigation request (e.g., /dashboard)",
            "Inspect the response headers",
            "Verify X-Content-Type-Options: nosniff is present",
            "Verify X-Frame-Options: DENY is present",
            "Verify X-XSS-Protection: 1; mode=block is present",
            "Also check API responses: curl -v http://localhost:8000/health",
            "Verify the API also returns security headers",
            "Verify no sensitive information in headers (no server version, no powered-by)",
            "(Production only) Verify Strict-Transport-Security header is present",
        ],
        "All security headers are present on both web and API responses. "
        "No information leakage in headers.",
    )

    sid = "S-16"
    scenarios.append((sid, "Input Validation"))
    add_scenario(
        doc,
        sid,
        "Input Validation: File Upload, SQL Injection, XSS",
        "All user inputs must be validated and sanitized to prevent attacks.",
        [
            "Navigate to /upload (resume upload page)",
            "Attempt to upload a .exe file",
            "Verify the upload is rejected with a clear error message (not a server crash)",
            "Attempt to upload a .bat file",
            "Verify the upload is rejected",
            "Navigate to /login",
            "Enter SQL injection in the email field: ' OR 1=1--",
            "Enter any password and click login",
            "Verify the response is a normal 'Invalid credentials' error "
            "(not a 500 server error)",
            "Verify no database error or stack trace is exposed",
            "Log in with a valid account, navigate to /profile",
            "Edit a text field (e.g., summary) and enter: <script>alert('xss')</script>",
            "Save the profile",
            "Refresh the page",
            "Verify the script tag is NOT executed (no alert popup)",
            "Verify the content is either escaped or stripped",
        ],
        "Dangerous file types are rejected at upload. SQL injection attempts are "
        "handled gracefully without server errors. XSS payloads are sanitized "
        "and never executed.",
    )

    sid = "S-17"
    scenarios.append((sid, "Auth Hardening"))
    add_scenario(
        doc,
        sid,
        "Authentication Hardening & Access Control",
        "Every protected endpoint must enforce authentication consistently.",
        [
            "Clear all cookies and auth tokens (use incognito mode)",
            "Attempt to access each protected API endpoint without auth:",
            "  GET /api/profile -> expect 401",
            "  GET /api/matches -> expect 401",
            "  GET /api/dashboard -> expect 401",
            "  GET /api/billing/status -> expect 401",
            "  POST /api/sieve/chat -> expect 401",
            "  POST /api/tailor -> expect 401",
            "Attempt to access admin endpoints without admin_token:",
            "  GET /api/admin/observability/health -> expect 401 or 403",
            "  GET /api/admin/trust/queue -> expect 401 or 403",
            "Attempt to use a manipulated JWT token "
            "(change a character in a valid token):",
            "  GET /api/auth/me with Authorization: Bearer <tampered-token> -> expect 401",
            "Verify CORS: from a different origin (or using curl with Origin header), "
            "confirm that cross-origin requests are handled per policy",
        ],
        "All protected endpoints return 401 without valid auth. Admin endpoints "
        "require admin token. Tampered JWTs are rejected. CORS policy is enforced.",
    )

    doc.add_page_break()

    # ================================================================
    # SECTION 5: WORKER & QUEUE VALIDATION (Phase 7 gaps)
    # ================================================================
    doc.add_heading("Section 5: Worker & Queue Validation", level=1)
    doc.add_paragraph(
        "These scenarios supplement the Recruiter/Admin UAT (R-14) with "
        "hands-on worker processing verification from the candidate perspective."
    )

    sid = "S-18"
    scenarios.append((sid, "Resume Parse Worker Job"))
    add_scenario(
        doc,
        sid,
        "Verify Resume Parse Worker Job End-to-End",
        "Resume parsing must complete reliably via the async worker pipeline.",
        [
            "Open a terminal showing the RQ worker output (python -m rq.cli worker)",
            "Log in as a candidate and navigate to /upload",
            "Upload a new PDF resume",
            "Watch the worker terminal for the parse job appearing",
            "Verify the worker logs show: job received, parsing started, parsing completed",
            "Return to the web app and verify the profile populates with parsed data",
            "Check the admin queue endpoint: "
            "GET /api/admin/observability/queues?admin_token=dev-admin-token",
            "Verify no stuck failed jobs in the parse queue",
        ],
        "Parse job appears in worker within seconds of upload. Worker completes "
        "parsing successfully. Profile populates. No failed jobs in queue.",
    )

    sid = "S-19"
    scenarios.append((sid, "Tailor Worker Job"))
    add_scenario(
        doc,
        sid,
        "Verify Tailored Resume Worker Job End-to-End",
        "Tailored resume generation must complete reliably via the async worker.",
        [
            "Ensure the RQ worker terminal is visible",
            "Navigate to a match detail page",
            "Click 'Generate ATS Resume'",
            "Watch the worker terminal for the tailor job appearing",
            "Verify the worker logs show: job received, tailoring started, tailoring completed",
            "Return to the web app and verify the download link appears",
            "Download the DOCX and verify it opens correctly",
            "Check the admin queue endpoint for any failed jobs",
            "Verify the failed count is 0 (or only known/expected failures)",
        ],
        "Tailor job runs in worker and completes. DOCX is downloadable and well-formed. "
        "No unexpected failed jobs.",
    )

    sid = "S-20"
    scenarios.append((sid, "Failed Job Inspection"))
    add_scenario(
        doc,
        sid,
        "Inspect and Retry Failed Worker Jobs",
        "Admins must be able to diagnose and recover from worker failures.",
        [
            "GET /api/admin/observability/queues?admin_token=dev-admin-token",
            "Check the 'failed' count for each queue (parse, tailor, match, etc.)",
            "If any failed jobs exist:",
            "  GET /api/admin/observability/queues/<queue_name>/failed?"
            "admin_token=dev-admin-token",
            "  Review the traceback and error message for each failed job",
            "  Determine if the failure is a real bug or a transient issue",
            "  POST /api/admin/observability/queues/<queue_name>/retry-all?"
            "admin_token=dev-admin-token",
            "  Verify retried jobs appear back in the queue and are re-processed",
            "If no failed jobs: verify the failed count is 0 across all queues",
        ],
        "Failed jobs (if any) have full tracebacks viewable. Retry moves them back "
        "to the queue for reprocessing. No orphaned failed jobs remain.",
    )

    doc.add_page_break()

    # ================================================================
    # SUMMARY CHECKLIST
    # ================================================================
    add_checklist_table(doc, scenarios)

    # ================================================================
    # SIGN-OFF
    # ================================================================
    doc.add_paragraph()
    doc.add_heading("Sign-Off", level=1)

    doc.add_paragraph(
        "This supplement covers 20 scenarios across 5 sections. Together with the "
        "main UAT docs (27 Candidate + 20 Employer + 20 Recruiter/Admin = 67 "
        "scenarios), the complete UAT suite totals 87 scenarios covering all 53 "
        "steps from the PROMPT25 QA checklist."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["Tester Name:", "Test Date:", "Environment:", "Overall Result:"]
    for i, label in enumerate(labels):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    path = os.path.join(SCRIPT_DIR, "uat-supplement.docx")
    doc.save(path)
    print(f"Created: {path}")
    return scenarios


if __name__ == "__main__":
    s = build_supplement_doc()
    print(f"  -> {len(s)} supplement scenarios")

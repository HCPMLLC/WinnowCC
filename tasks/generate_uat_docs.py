"""Generate three UAT DOCX scripts for Winnow: Candidate, Employer, Recruiter/Admin."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from lxml import etree
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elem = etree.SubElement(tc_pr, qn("w:shd"))
    shading_elem.set(qn("w:val"), "clear")
    shading_elem.set(qn("w:color"), "auto")
    shading_elem.set(qn("w:fill"), color_hex)


def make_doc(title, subtitle):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Winnow Job Matching Platform")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  |  February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()
    return doc


def add_section_header(doc, text):
    doc.add_heading(text, level=1)


def add_subsection(doc, text):
    doc.add_heading(text, level=2)


def add_differentiator(doc, text):
    """Add a differentiator callout box."""
    p = doc.add_paragraph()
    run = p.add_run("WINNOW DIFFERENTIATOR: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x0A, 0x7E, 0x3F)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x0A, 0x7E, 0x3F)


def add_scenario(doc, scenario_id, title, benefit, steps, expected, differentiator=None):
    """Add a test scenario with pass/fail checkbox."""
    doc.add_heading(f"{scenario_id}: {title}", level=3)

    p = doc.add_paragraph()
    run = p.add_run("Benefit: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(benefit)
    run.font.size = Pt(9)
    run.font.italic = True

    doc.add_paragraph("Steps:", style="List Bullet")
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Expected Result: ")
    run.bold = True
    run = p.add_run(expected)

    if differentiator:
        add_differentiator(doc, differentiator)

    # Pass/Fail line
    p = doc.add_paragraph()
    run = p.add_run("Result:  ")
    run.bold = True
    run = p.add_run("[ ] PASS    [ ] FAIL    Tester: ___________    Date: ___________")
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run("Notes: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run("_" * 80)
    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer


def add_checklist_table(doc, scenarios):
    """Add a summary pass/fail checklist table."""
    doc.add_heading("Summary Checklist", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["ID", "Scenario", "Pass/Fail", "Notes"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, "1B3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for sid, title in scenarios:
        row = table.add_row()
        row.cells[0].text = sid
        row.cells[1].text = title
        row.cells[2].text = "[ ] P  [ ] F"
        row.cells[3].text = ""
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)


# ============================================================
# CANDIDATE UAT
# ============================================================
def build_candidate_doc():
    doc = make_doc(
        "User Acceptance Test Script",
        "Candidate (Job Seeker) Persona",
    )

    add_section_header(doc, "Persona Overview")
    doc.add_paragraph(
        "The Candidate is a job seeker who uploads their resume, builds a profile, "
        "receives AI-scored job matches, generates tailored resumes and cover letters, "
        "tracks applications, and interacts with the Sieve AI assistant. Candidates care "
        "about finding the right job faster with less noise, understanding their chances, "
        "and presenting themselves optimally for each opportunity."
    )

    add_section_header(doc, "Key Benefits & Competitive Differentiators")
    benefits = [
        ("Interview Probability Score (IPS)", "Industry-first metric combining resume fit, cover letter quality, application timing, and referrals to predict interview likelihood."),
        ("Winnowing Philosophy", "Only shows matches scoring 50+ from the last 15 days. Fewer, higher-quality matches vs. the firehose approach of Indeed/LinkedIn."),
        ("AI-Grounded Resume Tailoring", "Every modification is grounded in actual profile data. Hallucination detection ensures nothing is fabricated."),
        ("Sieve AI with Proactive Triggers", "Context-aware AI assistant that proactively nudges candidates about incomplete profiles, stale applications, and high-value matches."),
        ("Semantic Job Search", "Natural language queries powered by pgvector embeddings (e.g., 'machine learning jobs in fintech')."),
        ("Skill Gap Analysis", "For every match, see exactly which skills you have and which you're missing vs. the job requirements."),
        ("GDPR-Compliant Data Export", "Full ZIP download including original resume files, not just metadata."),
        ("Trust & Transparency", "Automated trust scoring protects the ecosystem. Candidates know they're competing in a clean marketplace."),
    ]
    for title, desc in benefits:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_section_header(doc, "Prerequisites")
    prereqs = [
        "Development environment running: Docker (Postgres + Redis), API, Worker, Web",
        "Test candidate account: candidate@test.com / TestPass123!",
        "Test resume files: one PDF and one DOCX with realistic content",
        "At least 20 active jobs in the database (run job ingestion or seed script)",
        "Stripe test mode configured (for billing scenarios)",
        "OAuth provider configured (for social login scenario)",
    ]
    for p_text in prereqs:
        doc.add_paragraph(p_text, style="List Bullet")

    doc.add_page_break()

    # --- SCENARIOS ---
    scenarios = []

    add_section_header(doc, "1. Onboarding & Authentication")

    sid = "C-01"
    scenarios.append((sid, "Email Registration"))
    add_scenario(doc, sid, "Email Registration & Onboarding",
        "New candidates can create an account and set job preferences in under 2 minutes.",
        [
            "Navigate to /signup",
            "Enter email: newcandidate@test.com, password: SecurePass1!, confirm password",
            "Select role: 'Candidate'",
            "Click 'Create Account'",
            "Verify redirect to /onboarding",
            "Step 1: Enter target roles (e.g., 'Software Engineer, Backend Developer'), location 'San Francisco, CA', work mode 'Remote', salary range $100k-$150k, employment type 'Full-time'",
            "Click 'Next'",
            "Step 2: Accept Terms of Service, MJASS consent, Data processing consent",
            "Select application mode: 'Review required'",
            "Click 'Complete Onboarding'",
        ],
        "User lands on /dashboard with personalized greeting. Onboarding status marked complete.",
        "Application mode choice (review vs auto-apply) lets candidates control their job search automation level, unlike Indeed/LinkedIn which offer no such control."
    )

    sid = "C-02"
    scenarios.append((sid, "OAuth Social Login"))
    add_scenario(doc, sid, "OAuth Social Login (Google/LinkedIn/GitHub)",
        "Candidates can sign up/in with a single click using existing social accounts.",
        [
            "Navigate to /login",
            "Click 'Continue with Google' (or LinkedIn/GitHub)",
            "Complete OAuth flow in provider popup",
            "Verify redirect back to Winnow",
        ],
        "User is authenticated and redirected to /dashboard (or /onboarding if first time). OAuth profile data pre-fills name/email.",
        "5 OAuth providers supported (LinkedIn, Google, GitHub, Microsoft, Apple) vs. typically 1-2 on competing platforms."
    )

    sid = "C-03"
    scenarios.append((sid, "Password Reset"))
    add_scenario(doc, sid, "Password Reset Flow",
        "Candidates can securely recover access if they forget their password.",
        [
            "Navigate to /login, click 'Forgot password?'",
            "Enter registered email address",
            "Click 'Send Reset Link'",
            "Open reset email (check test mailbox or API logs)",
            "Click reset link, enter new password",
            "Submit, verify redirect to /login",
            "Log in with new password",
        ],
        "Password is changed. Reset token expires after 30 minutes. Old token no longer works.",
    )

    add_section_header(doc, "2. Resume Upload & Parsing")

    sid = "C-04"
    scenarios.append((sid, "PDF Resume Upload"))
    add_scenario(doc, sid, "Upload and Parse PDF Resume",
        "Candidates see their resume transformed into a structured profile in seconds.",
        [
            "Log in as candidate, navigate to /upload",
            "Click 'Upload Resume' and select a PDF file (<10MB)",
            "Verify upload progress indicator appears",
            "Wait for parsing to complete (status polling every 20 seconds)",
            "Verify redirect to /profile when complete",
            "Review parsed data: name, email, phone, location, experience entries, education, skills, certifications",
        ],
        "Profile is populated with all extracted fields. Skills are deduplicated. Experience bullets include quantified accomplishments. Phone is formatted as (NNN) NNN-NNNN.",
        "AI parsing with grounding validation ensures no hallucinated data. LLM fallback chain (OpenAI -> Anthropic -> regex) guarantees parsing even if primary provider is down."
    )

    sid = "C-05"
    scenarios.append((sid, "DOCX Resume Upload"))
    add_scenario(doc, sid, "Upload and Parse DOCX Resume",
        "Candidates using Word format get the same quality parsing as PDF users.",
        [
            "Navigate to /upload",
            "Upload a .docx resume file",
            "Wait for parsing completion",
            "Verify all fields extracted correctly (compare to original document)",
        ],
        "DOCX is parsed with same fidelity as PDF. Complex layouts (tables, columns) are handled by intelligent disambiguation.",
        "Resume disambiguation service handles complex/ambiguous resume formats that trip up competitors' parsers."
    )

    sid = "C-06"
    scenarios.append((sid, "Trust Scoring on Upload"))
    add_scenario(doc, sid, "Trust Score Computed Transparently",
        "Candidates benefit from a fraud-free marketplace where all participants are vetted.",
        [
            "Upload a legitimate resume with real-looking data",
            "Verify upload succeeds and profile is created",
            "Navigate to /dashboard",
            "Confirm no quarantine banner appears",
            "(Admin view: verify trust score computed with 4 buckets: identity, resume quality, online presence, abuse detection)",
        ],
        "Legitimate resumes pass trust gate automatically. Trust score > threshold = auto-approved. Candidates never see their numerical score, only whether they're approved.",
        "Automated 4-bucket trust scoring (identity 25pts, resume quality 20pts, online presence 25pts, abuse detection 30pts) protects all candidates from competing against fraudulent profiles."
    )

    add_section_header(doc, "3. Profile Management")

    sid = "C-07"
    scenarios.append((sid, "Profile Completeness"))
    add_scenario(doc, sid, "Profile Completeness Score & Recommendations",
        "Candidates know exactly what to improve for better match quality.",
        [
            "Navigate to /profile",
            "Note the completeness score (e.g., 65%)",
            "Review deficiency recommendations (e.g., 'Add certifications', 'Specify work authorization')",
            "Fill in a missing field (e.g., work authorization)",
            "Save profile and verify score increases",
            "Note the color progression: red (<50%) -> amber (50-79%) -> green (80%+)",
        ],
        "Completeness score updates immediately. Recommendations list shrinks as fields are filled. Visual progress ring reflects the new score.",
        "Weighted 100-point completeness scoring (Basics 20, Experience 40, Education 15, Skills 15, Preferences 10) gives candidates targeted improvement guidance, unlike platforms that just say 'complete your profile'."
    )

    sid = "C-08"
    scenarios.append((sid, "Skill Categorization"))
    add_scenario(doc, sid, "Skill Categorization for Better Matching",
        "Candidates can categorize skills to improve ATS parsing and match accuracy.",
        [
            "Navigate to /profile, scroll to Skills section",
            "Click 'Categorize Skills'",
            "Drag or assign skills to categories: Core Technical, Environmental/Adjacent, Leadership/Soft",
            "Save categorization",
            "Verify categories are preserved (navigate away and return)",
        ],
        "Skills are tagged with categories. Tailored resumes use these categories to group skills for optimal ATS parsing.",
        "Skill categorization improves ATS compatibility by organizing skills into sections that automated resume screeners expect, something no other matching platform offers."
    )

    sid = "C-09"
    scenarios.append((sid, "Profile Versioning"))
    add_scenario(doc, sid, "Profile Versioning Preserves History",
        "Candidates can update their profile without fear of losing previous versions.",
        [
            "View current profile, note version number",
            "Make a change (e.g., add a new skill)",
            "Save profile",
            "Verify version number incremented",
            "Generate a tailored resume",
            "Verify tailored resume references the correct profile version",
        ],
        "Each save creates a new version. Tailored resumes are linked to the specific profile version used, ensuring consistency.",
    )

    add_section_header(doc, "4. Job Matching & Discovery")

    sid = "C-10"
    scenarios.append((sid, "View Job Matches"))
    add_scenario(doc, sid, "View Matched Jobs with Quality Scores",
        "Candidates see only high-quality, relevant, fresh job matches -- not noise.",
        [
            "Navigate to /matches",
            "Verify matches are displayed with: job title, company, location, IPS score, match score",
            "Confirm all matches have IPS >= 50 (winnowing threshold)",
            "Confirm all matches are from jobs posted within the last 15 days",
            "Check that duplicate jobs (same role at same company from different boards) appear only once",
            "Note the source badges showing which board(s) each job was found on",
        ],
        "Match list shows only quality matches. Each card displays IPS score with color coding (green 80+, amber 60-79, blue 50-59). Stale and duplicate jobs are filtered out.",
        "Winnowing philosophy: only matches scoring 50+ from the last 15 days. Competitors show everything, overwhelming candidates with noise. Winnow shows fewer, higher-quality matches by design."
    )

    sid = "C-11"
    scenarios.append((sid, "Interview Probability Score"))
    add_scenario(doc, sid, "Interview Probability Score (IPS) Breakdown",
        "Candidates understand their actual chances and can take action to improve them.",
        [
            "Click on a match to open the detail panel",
            "Locate the IPS score and click to expand the breakdown",
            "Review the 4 components: Resume Score (35%), Cover Letter Score (20%), Application Logistics Score (25%), Referral Bonus (20%)",
            "Note specific improvement suggestions (e.g., 'Apply within first 10 days for +10 pts')",
            "Toggle the referral flag on a match and observe IPS increase",
        ],
        "IPS breakdown shows each component with points earned and max possible. Actionable improvement tips are specific and data-driven.",
        "Interview Probability Score (IPS) is an industry-first metric. No other platform quantifies interview likelihood using resume fit + cover letter quality + timing + referral status."
    )

    sid = "C-12"
    scenarios.append((sid, "Skill Gap Analysis"))
    add_scenario(doc, sid, "Skill Gap Analysis per Match",
        "Candidates see exactly which skills to highlight and which to develop.",
        [
            "Open a match detail panel",
            "Review 'Matched Skills' section (skills you have that the job requires)",
            "Review 'Missing Skills' section (required skills you lack)",
            "Review 'Skills to Highlight' section (your matched skills to emphasize on resume)",
            "Verify skill counts are accurate against your profile and the job requirements",
        ],
        "Clear visual breakdown of skill alignment. Matched skills shown in green, missing in red. Evidence references link matched skills to specific resume bullets.",
        "Evidence-based skill matching links each skill to the specific experience bullet where it was demonstrated, giving candidates (and employers) proof, not just assertions."
    )

    sid = "C-13"
    scenarios.append((sid, "Semantic Job Search"))
    add_scenario(doc, sid, "Natural Language Job Search",
        "Candidates can search for jobs using conversational queries, not just keywords.",
        [
            "Navigate to /matches",
            "In the search bar, type a natural language query: 'machine learning jobs at startups in healthcare'",
            "Submit the search",
            "Review results ranked by semantic similarity score",
            "Verify results are relevant to the query intent (not just keyword matches)",
        ],
        "Results include jobs that match the intent even if they don't contain the exact words. Similarity scores shown as percentages. Top 20 results returned with min 50% similarity.",
        "pgvector-powered semantic search understands intent, not just keywords. 'ML jobs in healthcare' finds 'Data Scientist at a health tech startup' -- keyword search would miss this."
    )

    sid = "C-14"
    scenarios.append((sid, "Similar Jobs Discovery"))
    add_scenario(doc, sid, "Find Similar Jobs from Any Match",
        "Candidates can expand their search based on jobs they like.",
        [
            "Open a match detail panel",
            "Click 'Find Similar Jobs'",
            "Review the 10 most similar jobs returned",
            "Verify they share relevant characteristics (industry, skills, seniority)",
        ],
        "Similar jobs are ranked by embedding cosine similarity. Results include jobs the candidate might not have seen through standard matching.",
    )

    sid = "C-15"
    scenarios.append((sid, "Match Refresh"))
    add_scenario(doc, sid, "Refresh Matches to Get Latest Jobs",
        "Candidates get fresh matches on demand, not on a platform-dictated schedule.",
        [
            "Navigate to /matches",
            "Click 'Refresh Matches'",
            "Verify loading state appears",
            "Wait for refresh to complete",
            "Confirm new matches appear (or count updates)",
            "Check usage counter: free tier shows '1 of 3 refreshes used'",
        ],
        "Match list updated with latest jobs. Usage counter incremented. Free tier: 3 refreshes/billing cycle.",
    )

    add_section_header(doc, "5. Tailored Resumes & Cover Letters")

    sid = "C-16"
    scenarios.append((sid, "Generate Tailored Resume"))
    add_scenario(doc, sid, "Generate ATS-Optimized Tailored Resume",
        "Candidates get a custom resume for each job that maximizes ATS compatibility.",
        [
            "Open a high-IPS match in the detail panel",
            "Click 'Generate Tailored Resume'",
            "Wait for generation (async via worker, status polling)",
            "Click 'Download Resume' to get the DOCX file",
            "Open the DOCX and verify: skills reordered to match job requirements, bullets enhanced with job keywords, section ordering optimized",
        ],
        "Clean DOCX with professional formatting. No tables (ATS-friendly). Skills grouped by category. Bullets emphasize matched requirements. No fabricated experience.",
        "Evidence-grounded tailoring: every modification traces back to actual profile data. Hallucination detection ensures nothing is invented. Competing tools (Jobscan, Kickresume) don't verify factual accuracy."
    )

    sid = "C-17"
    scenarios.append((sid, "Generate Cover Letter"))
    add_scenario(doc, sid, "Generate Personalized Cover Letter",
        "Candidates get a tailored cover letter that scores well on the IPS cover letter component.",
        [
            "After generating a tailored resume (C-16), click 'Download Cover Letter'",
            "Open the DOCX and verify:",
            "  - Addresses hiring manager by name (if available)",
            "  - Mentions the company name at least 2 times",
            "  - Uses enthusiasm keywords (excited, passionate, eager)",
            "  - Connects specific candidate experience to job requirements",
            "  - Length is 250-400 words (ATS optimal range)",
            "Verify no placeholder text (e.g., [Company Name]) remains",
        ],
        "Professional cover letter with no placeholders, accurate company/role references, appropriate tone, and ATS-optimal length.",
        "Cover letter quality directly feeds into the IPS calculation. Candidates can see their IPS cover letter component score increase after generating a tailored letter."
    )

    sid = "C-18"
    scenarios.append((sid, "Change Log & Grounding"))
    add_scenario(doc, sid, "Review Tailoring Change Log & Grounding Verification",
        "Candidates can trust that nothing was fabricated in their tailored documents.",
        [
            "Navigate to a generated tailored resume",
            "Click 'View Changes' or 'Change Log'",
            "Review: summary of modifications, per-section changes with reasons",
            "Verify keyword alignment: before/after match rate shown",
            "Check grounding badges: employers verified, titles verified, dates verified, education verified",
            "Confirm hallucination count is 0",
        ],
        "Change log shows exactly what was modified and why. Grounding verification confirms all facts are from the original profile. Zero hallucinations detected.",
        "Grounding validation with hallucination detection is unique to Winnow. The system verifies every employer name, job title, date, and credential against the original resume before finalizing."
    )

    add_section_header(doc, "6. Application Tracking")

    sid = "C-19"
    scenarios.append((sid, "Track Applications"))
    add_scenario(doc, sid, "Track Application Status Across Pipeline",
        "Candidates have a single view of all applications and their progress.",
        [
            "Navigate to /applications",
            "Verify pipeline columns: Saved, Applied, Interviewing, Offer, Rejected",
            "From /matches, change a match status to 'Applied'",
            "Return to /applications, verify it appears in the 'Applied' column",
            "Change status to 'Interviewing', verify card moves to correct column",
            "Review funnel visualization at the top showing conversion rates",
        ],
        "Application cards move between columns as status changes. Funnel chart shows conversion percentages at each stage. Status syncs between /matches and /applications pages.",
        "Kanban + funnel hybrid view gives candidates both task management and analytics perspectives. Competing platforms offer one or the other, not both."
    )

    sid = "C-20"
    scenarios.append((sid, "Application Notes"))
    add_scenario(doc, sid, "Add Notes to Job Applications",
        "Candidates can track details, follow-ups, and contacts for each application.",
        [
            "Open a match that has status 'Applied'",
            "Enter notes: 'Spoke with recruiter Jane on 2/10. Follow up next week.'",
            "Save notes",
            "Navigate away and return to the match",
            "Verify notes are preserved",
        ],
        "Notes persist across sessions. Free-text field allows any content.",
    )

    add_section_header(doc, "7. Sieve AI Assistant")

    sid = "C-21"
    scenarios.append((sid, "Sieve Proactive Triggers"))
    add_scenario(doc, sid, "Receive Proactive AI Nudges",
        "Candidates don't miss important actions -- the AI reminds them.",
        [
            "Log in with an account that has: incomplete profile (<70%), saved jobs older than 5 days, high-scoring unreviewed matches",
            "Navigate to /dashboard",
            "Observe Sieve widget showing proactive trigger cards",
            "Verify triggers include: 'Complete your profile', 'Review stale saved jobs', 'Check high-scoring matches'",
            "Dismiss a trigger and verify it doesn't reappear",
        ],
        "Sieve surfaces 1-3 relevant nudges based on the candidate's actual state. Triggers are dismissible and don't repeat.",
        "7-type proactive trigger system (profile completeness, new matches, stale saved jobs, high-scoring unreviewed, no tailored resumes, usage limits, interview coaching) is context-aware and data-driven -- not generic tips."
    )

    sid = "C-22"
    scenarios.append((sid, "Sieve Chat Interaction"))
    add_scenario(doc, sid, "Chat with Sieve AI About Job Search",
        "Candidates get personalized, context-aware career guidance.",
        [
            "Open the Sieve chat widget",
            "Ask: 'Which of my matches should I apply to first?'",
            "Verify Sieve references your actual matches, IPS scores, and profile data",
            "Ask: 'How can I improve my profile?'",
            "Verify Sieve references your actual completeness score and specific missing fields",
            "Ask: 'Help me prepare for an interview at [company from your matches]'",
            "Verify response includes relevant details from the job description",
        ],
        "Sieve responds with personalized advice using actual profile, match, and application data. Conversation history is maintained. Rate limited to 10 messages/minute.",
        "Unlike generic chatbots, Sieve has full context: your profile, all matches, application statuses, tailored resume history, and billing status. It gives specific, actionable advice rather than generic career tips."
    )

    add_section_header(doc, "8. Billing & Subscription")

    sid = "C-23"
    scenarios.append((sid, "Free Tier Limits"))
    add_scenario(doc, sid, "Free Tier Usage Limits & Upgrade Prompts",
        "Free tier candidates can use core features while seeing clear upgrade paths.",
        [
            "Log in with a free tier account",
            "Navigate to /matches -- verify only 10 matches shown",
            "Navigate to /settings -- verify usage bars: X/3 refreshes, X/5 tailored resumes",
            "Use all 3 match refreshes",
            "Attempt a 4th refresh -- verify graceful upgrade prompt (not an error)",
            "Use all 5 tailor requests",
            "Attempt a 6th -- verify upgrade prompt with pricing",
        ],
        "Free tier limits are enforced softly. When limits are reached, candidates see a clear upgrade CTA with pricing, not a hard error.",
        "Soft limit enforcement with graceful upgrade CTAs. Candidates understand exactly what they get free and what Pro unlocks, unlike platforms that gate-keep behind confusing paywalls."
    )

    sid = "C-24"
    scenarios.append((sid, "Upgrade to Pro"))
    add_scenario(doc, sid, "Upgrade to Pro via Stripe Checkout",
        "Candidates can upgrade seamlessly without leaving the platform.",
        [
            "Click 'Upgrade to Pro' from any upgrade prompt (or /settings)",
            "Select billing cycle: Monthly or Annual",
            "Verify redirect to Stripe Checkout page",
            "Complete payment with test card (4242 4242 4242 4242)",
            "Verify redirect back to Winnow /settings?billing=success",
            "Verify success toast: 'Welcome to Winnow Pro!'",
            "Verify plan badge updated to 'Pro'",
            "Verify match limit removed, refresh limit removed, tailor limit increased",
        ],
        "Subscription activated immediately. All Pro features unlocked. Usage counters reflect new limits.",
    )

    add_section_header(doc, "9. Data Export & Privacy")

    sid = "C-25"
    scenarios.append((sid, "GDPR Data Export"))
    add_scenario(doc, sid, "Export All Personal Data (GDPR Compliance)",
        "Candidates own their data and can take it with them.",
        [
            "Navigate to /settings",
            "In the Data Export section, click 'Preview Export'",
            "Review counts: profile versions, resume documents, matches, tailored resumes",
            "Click 'Download Export'",
            "Open the ZIP file and verify contents:",
            "  - account.json (user data)",
            "  - profiles/ folder (all profile versions as JSON)",
            "  - resumes/ folder (original PDF/DOCX files)",
            "  - matches.json (all match records)",
            "  - tailored/ folder (generated DOCX files)",
            "  - trust.json (trust records and audit log)",
        ],
        "ZIP download contains all personal data in organized folders. Includes original binary files (resume PDFs, tailored DOCXs), not just metadata.",
        "Full file export (original uploads + generated documents) goes beyond competitors who only export metadata or CSV summaries."
    )

    sid = "C-26"
    scenarios.append((sid, "Account Deletion"))
    add_scenario(doc, sid, "Permanently Delete Account",
        "Candidates can fully remove their presence from the platform.",
        [
            "Navigate to /settings",
            "Click 'Delete Account' in the danger zone",
            "Read the warning about irreversible deletion",
            "Type 'DELETE MY ACCOUNT' in the confirmation field",
            "Click 'Permanently Delete'",
            "Verify redirect to landing page",
            "Attempt to log in with deleted account -- verify it fails",
        ],
        "All data cascade-deleted: user, profiles, resumes, matches, tailored resumes, trust records, conversations. Auth cookie cleared.",
    )

    add_section_header(doc, "10. Competitive Comparison")

    sid = "C-27"
    scenarios.append((sid, "Feature Comparison Page"))
    add_scenario(doc, sid, "Review Feature Comparison vs. Competitors",
        "Candidates can make an informed choice about which platform to use.",
        [
            "Navigate to /competitive (or scroll to #compare on landing page)",
            "Review comparison table: Winnow vs. Indeed, LinkedIn, ZipRecruiter, Glassdoor",
            "Verify IPS is marked as Winnow-only",
            "Verify tailored resume generation is marked as Winnow-only",
            "Verify Sieve AI is marked as Winnow-only",
        ],
        "Comparison table accurately reflects feature differentiation. IPS, AI tailoring, and Sieve are clearly marked as unique to Winnow.",
    )

    doc.add_page_break()
    add_checklist_table(doc, scenarios)

    # Sign-off
    doc.add_paragraph()
    doc.add_heading("Sign-Off", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["Tester Name:", "Test Date:", "Environment:", "Overall Result:"]
    for i, label in enumerate(labels):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    path = os.path.join(SCRIPT_DIR, "uat-candidate.docx")
    doc.save(path)
    print(f"Created: {path}")
    return scenarios


# ============================================================
# EMPLOYER UAT
# ============================================================
def build_employer_doc():
    doc = make_doc(
        "User Acceptance Test Script",
        "Employer Persona",
    )

    add_section_header(doc, "Persona Overview")
    doc.add_paragraph(
        "The Employer is a hiring manager or HR professional who posts jobs, searches "
        "for candidates, manages multi-board distribution, tracks analytics, and manages "
        "their subscription. Employers care about reaching qualified candidates quickly, "
        "distributing jobs efficiently across multiple boards from a single platform, and "
        "understanding which channels deliver the best ROI."
    )

    add_section_header(doc, "Key Benefits & Competitive Differentiators")
    benefits = [
        ("One-Click Multi-Board Distribution", "Publish once, distribute to Indeed, Google Jobs, and custom XML feeds simultaneously. Auto-sync changes, auto-remove on close."),
        ("AI-Powered Job Document Parsing", "Upload a .docx job description and Winnow extracts all structured fields automatically with confidence scoring."),
        ("Privacy-Respecting Candidate Search", "Search candidates who opt-in to visibility. Anonymous profiles protect PII while showing skills and experience."),
        ("Auto-Distribution Hooks", "Job status changes trigger automatic board actions: active = distribute, paused/closed = remove, edit = sync."),
        ("Tiered Access with Clear Limits", "Free, Starter ($99/mo), and Pro ($299/mo) tiers with transparent limits on jobs and candidate views."),
        ("Centralized Cross-Board Metrics", "See impressions, clicks, applications, and cost per board in one dashboard instead of logging into 5 different platforms."),
        ("Real-Time Distribution Status", "Per-board status tracking (pending, live, failed, removed) with error messages and event audit log."),
    ]
    for title, desc in benefits:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_section_header(doc, "Prerequisites")
    prereqs = [
        "Development environment running: Docker (Postgres + Redis), API, Worker, Web",
        "Test employer account: employer@test.com / TestPass123! (role: employer)",
        "Test .docx job description file with realistic content",
        "At least 5 candidate profiles in the database (mix of public/anonymous/private visibility)",
        "Stripe test mode configured for billing scenarios",
        "Redis worker running for async distribution jobs",
    ]
    for p_text in prereqs:
        doc.add_paragraph(p_text, style="List Bullet")

    doc.add_page_break()

    scenarios = []

    add_section_header(doc, "1. Employer Profile & Onboarding")

    sid = "E-01"
    scenarios.append((sid, "Employer Profile Creation"))
    add_scenario(doc, sid, "Create Employer Profile",
        "Employers can set up their company presence quickly.",
        [
            "Log in as employer user, navigate to /employer/onboarding (or auto-redirect)",
            "Enter: company name 'Acme Corp', industry 'Technology', company size '51-200'",
            "Enter company website, description, and billing email",
            "Click 'Create Profile'",
            "Verify redirect to /employer/dashboard",
        ],
        "Employer profile created. Dashboard shows initial metrics (all zeros). Subscription defaults to Free tier.",
    )

    sid = "E-02"
    scenarios.append((sid, "View Subscription Tier Limits"))
    add_scenario(doc, sid, "Understand Subscription Tier Limits",
        "Employers know exactly what their current plan allows.",
        [
            "Navigate to /employer/settings",
            "Review current tier badge (Free)",
            "Note limits: 1 active job, 10 candidate views/month",
            "Click 'Upgrade' to see Starter (5 jobs, 50 views) and Pro (unlimited, 200 views) options",
        ],
        "Tier limits are clearly displayed. Upgrade options show pricing and feature differences.",
        "Transparent tiered pricing with clear feature matrix. No hidden limits or surprise charges."
    )

    add_section_header(doc, "2. Job Posting & Management")

    sid = "E-03"
    scenarios.append((sid, "Create Job Manually"))
    add_scenario(doc, sid, "Create a Job Posting Manually",
        "Employers can post jobs with full control over every field.",
        [
            "Navigate to /employer/jobs/new",
            "Fill in: title 'Senior Backend Engineer', description (200+ chars), requirements",
            "Set: location 'San Francisco, CA', remote policy 'Hybrid', employment type 'Full-time'",
            "Set: salary range $150k-$200k USD, equity offered = Yes",
            "Set: department 'Engineering', category 'Software Development'",
            "Click 'Save as Draft'",
            "Verify job appears in /employer/jobs with status 'Draft'",
        ],
        "Job created in draft status. Not distributed to any boards. All fields saved correctly.",
    )

    sid = "E-04"
    scenarios.append((sid, "Upload .docx Job Description"))
    add_scenario(doc, sid, "Upload .docx and Auto-Parse Job Description",
        "Employers can turn existing Word documents into structured job postings instantly.",
        [
            "Navigate to /employer/jobs, click 'Upload Document'",
            "Select a .docx file containing a job description",
            "Wait for AI parsing to complete",
            "Review auto-extracted fields: title, description, requirements, salary, location",
            "Note the parsing confidence score",
            "Edit any fields that need correction",
            "Click 'Save as Draft'",
        ],
        "Job draft created with all fields extracted from the document. Confidence score indicates parsing quality. Employer can review and correct before publishing.",
        "AI-powered document parsing with confidence scoring. Competing platforms require manual data entry for every field."
    )

    sid = "E-05"
    scenarios.append((sid, "Publish Job (Draft to Active)"))
    add_scenario(doc, sid, "Publish Job and Trigger Auto-Distribution",
        "Publishing a job automatically distributes it to all connected boards.",
        [
            "Navigate to /employer/jobs, find a draft job",
            "Click 'Publish' (or change status to 'Active')",
            "Verify status changes to 'Active'",
            "Verify posted_at timestamp is set",
            "Check the API logs or worker output for 'Enqueued auto-distribution for job X'",
            "Navigate to the job detail page, scroll to Distribution section",
            "Verify distribution records are created (pending or live) for each connected board",
        ],
        "Job goes active. Auto-distribution is triggered asynchronously. Distribution records appear with per-board status.",
        "Auto-distribution on publish: job is pushed to all connected boards within seconds. No manual posting to each board individually."
    )

    sid = "E-06"
    scenarios.append((sid, "Edit Active Job (Trigger Sync)"))
    add_scenario(doc, sid, "Edit Active Job and Auto-Sync to Boards",
        "Changes propagate to all boards automatically, keeping listings consistent.",
        [
            "Open an active, distributed job",
            "Change the title or salary range",
            "Save the changes",
            "Check worker logs for 'Enqueued distribution update for job X'",
            "Verify distribution records show updated feed_payload",
        ],
        "Job changes are pushed to all boards where the job is live. Distribution events log the update.",
        "Automatic cross-board sync on content changes. Editing once in Winnow updates Indeed, Google Jobs, and all connected boards."
    )

    sid = "E-07"
    scenarios.append((sid, "Pause/Close Job (Trigger Removal)"))
    add_scenario(doc, sid, "Pause or Close Job and Auto-Remove from Boards",
        "Stale listings are removed instantly, preventing candidate frustration.",
        [
            "Open an active, distributed job",
            "Change status to 'Paused'",
            "Check worker logs for 'Enqueued auto-removal for job X'",
            "Navigate to job detail, verify distribution status changes to 'removed'",
            "Verify removed_at timestamps are set on distribution records",
        ],
        "Job removed from all boards within minutes. Candidates on external boards no longer see the listing.",
        "Auto-removal on pause/close solves the industry's stale listing problem. Current platforms leave expired jobs visible for 6-24 hours."
    )

    sid = "E-08"
    scenarios.append((sid, "Archive and Unarchive Jobs"))
    add_scenario(doc, sid, "Archive and Restore Job Postings",
        "Employers can manage their job history without permanent deletion.",
        [
            "Open a closed job, click 'Archive'",
            "Provide reason: 'Position filled'",
            "Verify job moves to Archived tab in /employer/jobs?archived=true",
            "Click 'Unarchive' on an archived job",
            "Verify it returns to Draft status",
        ],
        "Archived jobs are hidden from active list but preserved. Unarchiving restores to draft for re-use.",
    )

    sid = "E-09"
    scenarios.append((sid, "Tier Limit Enforcement"))
    add_scenario(doc, sid, "Job Posting Tier Limits Enforced",
        "Employers understand their limits and get clear upgrade paths.",
        [
            "As a Free tier employer, create 1 active/draft job (succeeds)",
            "Attempt to create a 2nd job",
            "Verify 403 error with message: 'Free tier allows 1 active job(s). Upgrade to post more.'",
            "Upgrade to Starter tier",
            "Verify can now create up to 5 jobs",
        ],
        "Tier limits enforced at job creation with clear error messages and upgrade suggestion.",
    )

    add_section_header(doc, "3. Multi-Board Distribution")

    sid = "E-10"
    scenarios.append((sid, "Add Board Connection"))
    add_scenario(doc, sid, "Connect to External Job Boards",
        "Employers set up board connections once and reuse them for all jobs.",
        [
            "Navigate to /employer/connections",
            "Click 'Add Connection'",
            "Select board type: 'Indeed'",
            "Enter: board name, API key, API secret",
            "Click 'Save'",
            "Verify connection appears in the list with 'Active' badge",
            "Add a second connection: 'Google for Jobs' (no credentials needed)",
            "Add a third: 'Custom XML Feed' with feed URL",
        ],
        "Three board connections created. Each shows type, name, status, and last sync info.",
        "Centralized board management: configure once, distribute to all. No logging into 5 different board admin panels."
    )

    sid = "E-11"
    scenarios.append((sid, "Test Board Credentials"))
    add_scenario(doc, sid, "Validate Board Connection Credentials",
        "Employers can verify their setup before distributing jobs.",
        [
            "Navigate to /employer/connections",
            "Click 'Test' on the Indeed connection",
            "Wait for credential validation to complete",
            "Verify result shows success or failure with specific error message",
        ],
        "Test endpoint validates credentials against the board's API and reports success/failure with details.",
    )

    sid = "E-12"
    scenarios.append((sid, "Distribution Status & Metrics"))
    add_scenario(doc, sid, "View Per-Board Distribution Status and Metrics",
        "Employers see a unified view of how their job performs across all boards.",
        [
            "Navigate to a distributed job's detail page",
            "Scroll to 'Board Distribution' section",
            "Verify each board shows: status badge (live/pending/failed), external job ID",
            "Verify metrics per board: impressions, clicks, applications, cost spent",
            "Click 'Sync Metrics' to pull latest data from all boards",
            "Verify metrics update after sync",
        ],
        "Per-board status and metrics in one view. Color-coded badges for quick scanning. Metrics sync on demand.",
        "Centralized cross-board visibility: see impressions, clicks, applications, and cost from Indeed, Google Jobs, and all boards in a single dashboard. No other platform provides this."
    )

    add_section_header(doc, "4. Candidate Search & Management")

    sid = "E-13"
    scenarios.append((sid, "Search Candidates by Skills"))
    add_scenario(doc, sid, "Search for Candidates by Skills, Location, and Title",
        "Employers find qualified, opt-in candidates without sifting through inactive profiles.",
        [
            "Navigate to /employer/candidates",
            "Enter skills filter: 'Python, React'",
            "Enter location filter: 'San Francisco'",
            "Click 'Search'",
            "Verify results only show candidates with open_to_opportunities = true",
            "Verify anonymous profiles show 'Candidate #ID' instead of real name",
            "Verify public profiles show full name and headline",
            "Check pagination: click 'Next' if results exceed page size",
        ],
        "Results show only opt-in candidates. Anonymous candidates display skills and experience without PII. Pagination works correctly.",
        "Privacy-respecting candidate search: only shows candidates who actively opted in. Anonymous profiles protect identity while showing qualifications."
    )

    sid = "E-14"
    scenarios.append((sid, "View Candidate Profile"))
    add_scenario(doc, sid, "View Candidate Profile (Counts as View)",
        "Employers get full candidate details while the platform tracks view usage.",
        [
            "Click on a candidate from search results",
            "Review: full profile JSON, experience, skills, education, years of experience",
            "Navigate to /employer/settings or /employer/dashboard",
            "Verify 'Candidate views this month' incremented by 1",
            "For anonymous candidate: verify PII is redacted but skills/experience visible",
        ],
        "Full profile displayed for public candidates. Anonymous candidates show skills/experience without PII. View counted toward monthly limit.",
    )

    sid = "E-15"
    scenarios.append((sid, "Candidate View Limit"))
    add_scenario(doc, sid, "Hit Candidate View Monthly Limit",
        "Employers get clear feedback when approaching and hitting limits.",
        [
            "As a Free tier employer (10 views/month), view 10 candidate profiles",
            "Attempt to view an 11th candidate",
            "Verify 403 error: 'Free tier allows 10 candidate views/month. Upgrade for more.'",
        ],
        "View limit enforced with clear error and upgrade suggestion.",
    )

    sid = "E-16"
    scenarios.append((sid, "Save & Manage Candidates"))
    add_scenario(doc, sid, "Save Candidates with Notes",
        "Employers can build a shortlist with context for each candidate.",
        [
            "From candidate search, click 'Save' on a candidate",
            "Add notes: 'Strong Python background, good culture fit'",
            "Navigate to /employer/candidates/saved",
            "Verify saved candidate appears with notes",
            "Edit notes to add: 'Schedule phone screen next week'",
            "Save changes, verify notes updated",
            "Remove a saved candidate, verify removed from list",
        ],
        "Saved candidates list with notes persists across sessions. CRUD operations work correctly.",
    )

    add_section_header(doc, "5. Analytics & Dashboard")

    sid = "E-17"
    scenarios.append((sid, "Analytics Dashboard"))
    add_scenario(doc, sid, "View Employer Analytics Dashboard",
        "Employers get real-time visibility into their hiring pipeline.",
        [
            "Navigate to /employer/dashboard",
            "Verify stats cards: active jobs count, total job views, total applications",
            "Verify candidate views this month with progress bar (X / limit)",
            "Verify saved candidates count",
            "Verify subscription tier badge",
            "Create a new job and publish it, return to dashboard",
            "Verify active jobs count incremented",
        ],
        "Dashboard shows real-time aggregated metrics. Numbers update immediately after actions.",
        "Real-time aggregated metrics across all jobs and boards in one view."
    )

    add_section_header(doc, "6. Billing & Subscription")

    sid = "E-18"
    scenarios.append((sid, "Stripe Checkout Upgrade"))
    add_scenario(doc, sid, "Upgrade Subscription via Stripe",
        "Employers can upgrade seamlessly with standard payment processing.",
        [
            "Navigate to /employer/settings",
            "Click 'Upgrade to Starter' or 'Upgrade to Pro'",
            "Verify redirect to Stripe Checkout",
            "Complete payment with test card (4242 4242 4242 4242)",
            "Verify redirect back to /employer/settings with success message",
            "Verify tier badge updated",
            "Verify new limits in effect (e.g., can now create 5 jobs on Starter)",
        ],
        "Subscription activated immediately. Tier limits updated. Stripe Customer Portal available for management.",
    )

    sid = "E-19"
    scenarios.append((sid, "Manage Subscription"))
    add_scenario(doc, sid, "Manage Subscription via Customer Portal",
        "Employers can self-serve their billing without contacting support.",
        [
            "Click 'Manage Subscription' in /employer/settings",
            "Verify redirect to Stripe Customer Portal",
            "Review: current plan, payment method, invoices, cancel option",
            "Return to Winnow",
        ],
        "Stripe Customer Portal opens in new tab/redirect. All subscription management available self-service.",
        "Full self-service subscription management via Stripe. No manual invoicing or sales calls required."
    )

    add_section_header(doc, "7. Job Lifecycle End-to-End")

    sid = "E-20"
    scenarios.append((sid, "Full Job Lifecycle"))
    add_scenario(doc, sid, "Complete Job Lifecycle: Draft to Archive",
        "The full employer workflow works end-to-end without gaps.",
        [
            "Create job as Draft (E-03)",
            "Publish to Active -- verify auto-distribution triggered (E-05)",
            "Edit title/salary -- verify auto-sync to boards (E-06)",
            "Pause job -- verify auto-removal from boards (E-07)",
            "Re-activate job -- verify re-distribution triggered",
            "Close job -- verify final removal from boards",
            "Archive job with reason 'Position filled' (E-08)",
            "Verify complete distribution event log on job detail page",
        ],
        "All status transitions trigger correct distribution actions. Event audit log captures every action with timestamps.",
        "Automated lifecycle management: status transitions drive distribution actions without manual intervention. This eliminates the #1 employer complaint about job boards."
    )

    doc.add_page_break()
    add_checklist_table(doc, scenarios)

    doc.add_paragraph()
    doc.add_heading("Sign-Off", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["Tester Name:", "Test Date:", "Environment:", "Overall Result:"]
    for i, label in enumerate(labels):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    path = os.path.join(SCRIPT_DIR, "uat-employer.docx")
    doc.save(path)
    print(f"Created: {path}")
    return scenarios


# ============================================================
# RECRUITER / ADMIN UAT
# ============================================================
def build_admin_doc():
    doc = make_doc(
        "User Acceptance Test Script",
        "Recruiter / Platform Admin Persona",
    )

    add_section_header(doc, "Persona Overview")
    doc.add_paragraph(
        "The Recruiter/Admin is a platform operator responsible for maintaining marketplace "
        "quality, managing trust and compliance, monitoring system health, and ensuring job "
        "data integrity. This persona may be an internal Winnow team member or a power-user "
        "recruiter with elevated access. They care about a clean, trustworthy marketplace, "
        "efficient operations, and proactive quality control."
    )

    add_section_header(doc, "Key Benefits & Competitive Differentiators")
    benefits = [
        ("Automated 4-Bucket Trust Scoring", "Identity, resume quality, online presence, and abuse detection scored automatically. Manual overrides for edge cases. Full audit trail."),
        ("14-Signal Job Fraud Detection", "Scam phrases, personal info requests, fee requirements, salary anomalies, duplicate detection, and more. Each signal has a severity and point value."),
        ("Purgeable Profile Scanner", "Auto-detects test profiles (regex on name/email) and inactive accounts (no onboarding, no resume, >7 days old) for batch cleanup."),
        ("Safe Candidate Merge", "Merge duplicate accounts while preserving all data: profiles, resumes, matches, tailored resumes, trust records."),
        ("Full System Observability", "Health probes, queue depth monitoring, failed job retry, structured JSON logging with PII redaction."),
        ("Multi-Source Job Ingestion", "13+ board sources (Remotive, The Muse, Greenhouse, Lever, USAJobs, etc.) with content hash dedup and freshness filtering."),
        ("Automated Security Posture Check", "One-click audit of AUTH_SECRET strength, environment mode, CORS, Stripe key mode, database encryption, and storage config."),
    ]
    for title, desc in benefits:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(10)

    add_section_header(doc, "Prerequisites")
    prereqs = [
        "Development environment running: Docker (Postgres + Redis), API, Worker, Web",
        "Admin account with X-Admin-Token header configured (ADMIN_TOKEN env var)",
        "Test candidates: 1 clean, 2 soft quarantine, 1 hard quarantine, 1 duplicate email",
        "Test jobs: 3 clean, 3 fraud score 40-59, 4 fraud score 60+",
        "At least 2 purgeable profiles: 1 test profile (test@example.com), 1 inactive (>7 days, no onboarding)",
        "Worker running for async job processing",
    ]
    for p_text in prereqs:
        doc.add_paragraph(p_text, style="List Bullet")

    doc.add_page_break()

    scenarios = []

    add_section_header(doc, "1. Trust Queue Management")

    sid = "R-01"
    scenarios.append((sid, "View Trust Queue"))
    add_scenario(doc, sid, "Review Quarantined Candidates in Trust Queue",
        "Admins can quickly identify and triage candidates flagged by automated trust scoring.",
        [
            "Navigate to /admin/trust",
            "Verify quarantine queue shows candidates with status: soft_quarantine, hard_quarantine",
            "Review each candidate's trust score breakdown:",
            "  - Identity bucket (25 pts): name, email, location, work history",
            "  - Resume Quality bucket (20 pts): parse success, job entries, overlapping dates, keyword stuffing",
            "  - Online Presence bucket (25 pts): LinkedIn, GitHub, portfolio URLs",
            "  - Abuse Detection bucket (30 pts): duplicate hashes, frequent uploads",
            "Verify total score and reason codes are displayed",
        ],
        "Trust queue shows all quarantined candidates with detailed score breakdowns. Each bucket score is visible with contributing signals.",
        "Automated 4-bucket trust scoring replaces manual resume review. Most platforms rely on simple duplicate detection or no trust layer at all."
    )

    sid = "R-02"
    scenarios.append((sid, "Override Trust Status"))
    add_scenario(doc, sid, "Override Trust Status with Audit Trail",
        "Admins can resolve edge cases while maintaining a complete audit history.",
        [
            "Select a soft-quarantined candidate",
            "Click 'Approve' (set status to 'allowed')",
            "Add admin note: 'Verified identity via phone call'",
            "Save the override",
            "Select another candidate, click 'Hard Quarantine' with note: 'Fraudulent resume detected'",
            "Navigate to the audit log section",
            "Verify both overrides are logged with: admin user, timestamp, old status, new status, notes",
        ],
        "Trust status changed. Audit log captures every override with full context. Candidate can now access matching (if approved) or is blocked (if quarantined).",
        "Immutable audit trail for every trust decision. Essential for compliance and dispute resolution."
    )

    sid = "R-03"
    scenarios.append((sid, "Auto-Trust Returning Users"))
    add_scenario(doc, sid, "Auto-Trust for Previously Approved Users",
        "Returning users don't face repeated vetting friction.",
        [
            "Upload a resume for a previously-approved candidate (one approved resume already exists)",
            "Verify the new upload is auto-approved (no quarantine)",
            "Check trust audit log for auto-trust event",
        ],
        "Second resume upload bypasses manual review. Auto-trust event logged with reason 'previous approval'.",
        "Auto-trust for returning users reduces friction while maintaining security. New users go through full vetting; returning users get fast-tracked."
    )

    add_section_header(doc, "2. Job Quality & Fraud Detection")

    sid = "R-04"
    scenarios.append((sid, "Review Flagged Jobs"))
    add_scenario(doc, sid, "Review Jobs Flagged by Fraud Detection",
        "Admins can protect candidates from fraudulent job postings.",
        [
            "Navigate to /admin/job-quality",
            "Verify jobs with fraud score >= 40 are listed",
            "Review a high-scoring job (60+):",
            "  - Fraud score with color coding (red >= 60, amber 40-59)",
            "  - Red flag details: signal name, severity (high/medium/low), point value",
            "  - Posting quality score (0-100)",
            "  - Full job description preview",
            "Review a medium-scoring job (40-59) for comparison",
        ],
        "Flagged jobs displayed with detailed fraud analysis. Each red flag shows specific evidence and severity.",
        "14-signal fraud detection: scam phrases (20pts), no company (15pts), personal info request (20pts), fee required (25pts), salary anomaly (10pts), urgency language (8pts), vague title (8pts), duplicate posting (5pts), and more."
    )

    sid = "R-05"
    scenarios.append((sid, "Override Fraud Status"))
    add_scenario(doc, sid, "Mark Job as Fraudulent or Legitimate",
        "Admins make final calls on borderline cases to maintain marketplace quality.",
        [
            "Select a flagged job",
            "Click 'Mark as Fraudulent'",
            "Verify job's is_active is set to false (candidates no longer see it)",
            "Select another flagged job that appears legitimate",
            "Click 'Mark as Legitimate'",
            "Verify job remains active with fraud override flag set",
        ],
        "Fraudulent jobs are deactivated immediately. Legitimate jobs are cleared. Both actions are logged.",
    )

    sid = "R-06"
    scenarios.append((sid, "Batch Reparse Jobs"))
    add_scenario(doc, sid, "Batch Reparse All Jobs for Updated Fraud Detection",
        "Admins can re-evaluate the entire job database when fraud rules are updated.",
        [
            "Navigate to /admin/job-quality",
            "Click 'Reparse All Jobs'",
            "Verify background job is enqueued (check worker logs)",
            "Wait for reparse to complete",
            "Verify fraud scores are recalculated across all jobs",
            "Check for newly flagged jobs that weren't flagged before",
        ],
        "All jobs re-evaluated with latest fraud detection rules. New flags surface previously missed issues.",
    )

    add_section_header(doc, "3. Candidate Management")

    sid = "R-07"
    scenarios.append((sid, "Search and Browse Candidates"))
    add_scenario(doc, sid, "Search Candidates by Name, Email, Title, Location",
        "Admins can quickly locate any candidate in the system.",
        [
            "Navigate to /admin/candidates",
            "Verify sortable table with: name, title, city, state, email, trust status, years exp",
            "Search by name: enter partial name, verify filtered results",
            "Search by email: enter email address, verify exact match",
            "Search by title: enter job title keyword",
            "Search by location: enter city or state",
            "Verify duplicate emails are highlighted in amber",
        ],
        "Table filters in real-time. Duplicate detection highlights potential issues. Trust status badges are color-coded.",
    )

    sid = "R-08"
    scenarios.append((sid, "Preview Resume"))
    add_scenario(doc, sid, "Preview Resumes Inline (PDF and DOCX)",
        "Admins can review resumes without downloading files.",
        [
            "Click on a candidate with a PDF resume",
            "Verify PDF renders inline in the browser",
            "Click on a candidate with a DOCX resume",
            "Verify DOCX is converted via mammoth.js and displayed as HTML",
            "Click 'Download' to get the original file",
        ],
        "Both PDF and DOCX resumes viewable inline. Original files downloadable.",
    )

    sid = "R-09"
    scenarios.append((sid, "Merge Duplicate Candidates"))
    add_scenario(doc, sid, "Merge Duplicate Candidate Accounts Safely",
        "Admins can consolidate duplicate accounts without losing any data.",
        [
            "Identify two candidates with the same email (duplicate highlight)",
            "Select both candidates using checkboxes",
            "Click 'Merge Selected'",
            "Select the primary account to keep",
            "Confirm merge in the confirmation dialog",
            "Verify: all profiles, resumes, matches, tailored resumes, trust records are transferred to primary",
            "Verify duplicate account is deleted",
            "Verify primary account has all data from both accounts",
        ],
        "Merge completes without data loss. All associated records transferred to primary account. Duplicate removed.",
        "Safe merge with complete data transfer: profiles, resumes, matches, tailored documents, and trust records all preserved. No other platform offers non-destructive account merging."
    )

    sid = "R-10"
    scenarios.append((sid, "Bulk Delete with Cascade"))
    add_scenario(doc, sid, "Bulk Delete Candidates with Full Cascade",
        "Admins can remove accounts completely with no orphaned records.",
        [
            "Select 2 candidates using checkboxes",
            "Click 'Delete Selected'",
            "Confirm in the confirmation dialog",
            "Verify all associated data deleted: profiles, resumes, matches, tailored resumes, trust records, audit logs, sieve conversations",
            "Verify no orphaned records in any related table",
        ],
        "Complete cascade deletion. No orphaned foreign key records. Admin accounts are protected from accidental deletion.",
    )

    add_section_header(doc, "4. Profile & User Administration")

    sid = "R-11"
    scenarios.append((sid, "Purgeable Profile Scanner"))
    add_scenario(doc, sid, "Detect Test and Inactive Profiles for Cleanup",
        "Admins keep the database clean by identifying accounts that serve no purpose.",
        [
            "Navigate to /admin/profile",
            "Click 'Scan for Purgeable Profiles'",
            "Review results:",
            "  - Test profiles: matched by regex on name/email (e.g., test@example.com, 'Test User')",
            "  - Inactive profiles: no onboarding completed, no resume uploaded, created >7 days ago",
            "Select purgeable profiles",
            "Click 'Purge Selected'",
            "Confirm in modal",
            "Verify profiles are deleted with cascade",
        ],
        "Scanner identifies test and inactive accounts. Purge removes them cleanly.",
        "Auto-detection of test profiles (regex matching) and inactive accounts (behavioral signals) keeps the database clean without manual hunting."
    )

    sid = "R-12"
    scenarios.append((sid, "Profile Completeness Overview"))
    add_scenario(doc, sid, "View Profile Completeness Across All Users",
        "Admins can identify users who need engagement nudges.",
        [
            "Navigate to /admin/profile",
            "Verify completeness indicators: green (80%+), amber (50-79%), red (<50%)",
            "Verify onboarding status badges",
            "Sort by completeness to find incomplete profiles",
            "Click a user to view their detailed profile",
        ],
        "Admin has visibility into profile completeness across the user base. Can identify engagement opportunities.",
    )

    add_section_header(doc, "5. System Health & Observability")

    sid = "R-13"
    scenarios.append((sid, "Health Check"))
    add_scenario(doc, sid, "Verify System Health Probes",
        "Admins can confirm all infrastructure components are operational.",
        [
            "Call GET /health (or navigate to health endpoint)",
            "Verify response includes: API status, uptime",
            "Call GET /ready",
            "Verify response checks: database connectivity, Redis connectivity",
            "If any check fails, verify clear error message identifying the failing component",
        ],
        "Both probes return healthy status. Ready probe confirms DB and Redis connections.",
    )

    sid = "R-14"
    scenarios.append((sid, "Queue Monitoring"))
    add_scenario(doc, sid, "Monitor Worker Queue Stats and Failed Jobs",
        "Admins have full visibility into background job processing.",
        [
            "Call GET /api/admin/observability/queues (with admin token)",
            "Verify queue stats for each queue: queued, started, finished, failed, deferred, scheduled",
            "Check for any failed jobs",
            "If failed jobs exist: view tracebacks via /queues/{queue_name}/failed",
            "Retry failed jobs via /queues/{queue_name}/retry-all",
            "Verify retried jobs appear back in the queue",
        ],
        "Queue stats show real-time processing status. Failed jobs have full tracebacks. Retry moves them back to the queue.",
        "Full queue observability with traceback inspection and one-click retry. No need to SSH into servers or use external monitoring tools."
    )

    sid = "R-15"
    scenarios.append((sid, "Security Posture Check"))
    add_scenario(doc, sid, "Run Automated Security Posture Audit",
        "Admins can verify platform security configuration with one click.",
        [
            "Call GET /api/admin/security/check (with admin token)",
            "Review each check:",
            "  - AUTH_SECRET: PASS if >= 32 chars and not default",
            "  - Environment: WARN if dev mode in production",
            "  - CORS origins: PASS if properly restricted",
            "  - Stripe keys: WARN if test mode in production",
            "  - Database encryption: PASS if Cloud SQL auto-encrypted",
            "  - GCS bucket: configured vs local storage",
            "Verify each returns PASS, WARN, or FAIL with explanation",
        ],
        "Security report generated with per-check status. Any WARN or FAIL items are actionable.",
        "Automated security posture check validates 6 critical configuration areas. No manual checklist required."
    )

    add_section_header(doc, "6. Job Ingestion & Data Quality")

    sid = "R-16"
    scenarios.append((sid, "Trigger Job Ingestion"))
    add_scenario(doc, sid, "Manually Trigger Multi-Source Job Ingestion",
        "Admins can refresh the job database on demand from 13+ sources.",
        [
            "Call POST /api/admin/scheduler/trigger (with admin token)",
            "Verify background job is enqueued",
            "Monitor worker logs for ingestion progress",
            "Wait for completion",
            "Verify new jobs appear in the database",
            "Check for: deduplication (content hash), freshness filtering (only last 7 days), fraud scoring on new jobs",
        ],
        "Job ingestion runs across all configured sources. New jobs added with dedup. Fraud scores computed. Embeddings generated.",
        "13+ board sources (Remotive, The Muse, Greenhouse, Lever, USAJobs, Adzuna, JSearch, etc.) with content hash deduplication and automatic fraud scoring on ingest."
    )

    sid = "R-17"
    scenarios.append((sid, "View Ingestion History"))
    add_scenario(doc, sid, "Review Job Ingestion Run History",
        "Admins can audit ingestion runs for errors and volume trends.",
        [
            "Call GET /api/admin/scheduler/runs (with admin token)",
            "Verify last 20 runs listed with: run ID, status, jobs ingested, errors, timestamps",
            "Review any failed runs for error messages",
            "Verify successful runs show non-zero job counts",
        ],
        "Run history shows full audit trail. Failed runs include error details for debugging.",
    )

    add_section_header(doc, "7. Distribution Operations")

    sid = "R-18"
    scenarios.append((sid, "Scheduled Metrics Sync"))
    add_scenario(doc, sid, "Verify Scheduled Distribution Metrics Sync",
        "The platform automatically keeps distribution metrics current.",
        [
            "Verify the scheduled_sync_distribution_metrics function runs every 15 minutes",
            "Check worker logs for: 'Distribution metrics sync: X synced, Y errors'",
            "After sync, verify job distribution records have updated: impressions, clicks, applications, cost_spent",
            "Verify board connection last_sync_at timestamps are updated",
        ],
        "Metrics sync runs on schedule. All live distributions have current metrics. Board sync status updated.",
    )

    sid = "R-19"
    scenarios.append((sid, "Expired Job Auto-Cleanup"))
    add_scenario(doc, sid, "Auto-Archive Expired Jobs and Remove from Boards",
        "Stale jobs are cleaned up automatically without admin intervention.",
        [
            "Create an employer job with close_date in the past",
            "Set status to 'active'",
            "Run the scheduled_archive_expired_jobs function (or wait for scheduled execution)",
            "Verify job status changed to 'closed', archived = true, archived_reason = 'expired'",
            "Verify auto-removal from boards triggered (check for process_removal in logs)",
            "Verify distribution records show status = 'removed'",
        ],
        "Expired jobs auto-archived and removed from all boards. No manual intervention required. Candidates no longer see stale listings.",
        "Automatic expired job cleanup with board removal solves the industry's stale listing problem at the infrastructure level."
    )

    add_section_header(doc, "8. Sieve AI Administration")

    sid = "R-20"
    scenarios.append((sid, "Sieve Trigger System"))
    add_scenario(doc, sid, "Verify Sieve Proactive Trigger Engine",
        "The AI assistant proactively engages users based on their state.",
        [
            "Create a test user with: incomplete profile (<50%), saved jobs older than 5 days, high-scoring unreviewed matches",
            "Call POST /api/sieve/triggers for this user",
            "Verify triggers returned include:",
            "  - Profile completeness nudge (profile < 70%)",
            "  - Stale saved jobs alert (saved > 5 days, not applied)",
            "  - High-scoring match notification (80%+ score, no status set)",
            "Dismiss a trigger and verify it doesn't reappear on next call",
            "Verify rate limiting: 10 messages/minute per user",
        ],
        "Trigger engine correctly identifies user states and generates relevant nudges. Dismissed triggers don't repeat.",
        "7-type proactive trigger system: profile completeness, new matches, stale saved jobs, high-scoring unreviewed, no tailored resumes, usage limits, interview coaching. Context-aware, not generic."
    )

    doc.add_page_break()
    add_checklist_table(doc, scenarios)

    doc.add_paragraph()
    doc.add_heading("Sign-Off", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    labels = ["Tester Name:", "Test Date:", "Environment:", "Overall Result:"]
    for i, label in enumerate(labels):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    path = os.path.join(SCRIPT_DIR, "uat-recruiter-admin.docx")
    doc.save(path)
    print(f"Created: {path}")
    return scenarios


if __name__ == "__main__":
    c = build_candidate_doc()
    print(f"  -> {len(c)} candidate scenarios")
    e = build_employer_doc()
    print(f"  -> {len(e)} employer scenarios")
    a = build_admin_doc()
    print(f"  -> {len(a)} recruiter/admin scenarios")
    print(f"\nTotal: {len(c) + len(e) + len(a)} test scenarios across 3 documents")

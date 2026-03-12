"""Submittal Package Service — builds and sends candidate packages to clients.

Orchestrates AI brief generation, DOCX building, PDF merging, and email
delivery for recruiter candidate submittals.
"""

from __future__ import annotations

import logging
import tempfile
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.db.session import get_session_factory
from app.models.candidate_profile import CandidateProfile
from app.models.candidate_submission import CandidateSubmission
from app.models.recruiter_client import RecruiterClient
from app.models.recruiter_job import RecruiterJob
from app.models.recruiter_pipeline_candidate import RecruiterPipelineCandidate
from app.models.resume_document import ResumeDocument
from app.models.submittal_package import SubmittalPackage
from app.models.tailored_resume import TailoredResume

logger = logging.getLogger(__name__)

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x2B, 0x6C, 0xB0)
GRAY = RGBColor(0x66, 0x66, 0x66)


# ---------------------------------------------------------------------------
# Cover page DOCX
# ---------------------------------------------------------------------------

def _build_cover_page(
    job: RecruiterJob,
    client_name: str | None,
    recipient_name: str,
    candidates: list[dict],
    output_path: str,
) -> str:
    """Build a professional cover page DOCX for the submittal package."""
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run("CANDIDATE SUBMITTAL")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    # Divider line
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div.paragraph_format.space_after = Pt(20)
    drun = div.add_run("_" * 60)
    drun.font.color.rgb = BLUE
    drun.font.size = Pt(8)

    # Job details
    details = [
        ("Position", job.title),
        ("Solicitation #", job.job_id_external or "N/A"),
        ("Client", client_name or job.client_company_name or "N/A"),
        ("Prepared For", recipient_name),
        ("Date", datetime.now(UTC).strftime("%B %d, %Y")),
        ("Candidates", str(len(candidates))),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        lbl = p.add_run(f"{label}: ")
        lbl.bold = True
        lbl.font.name = "Calibri"
        lbl.font.size = Pt(12)
        lbl.font.color.rgb = NAVY
        val = p.add_run(value)
        val.font.name = "Calibri"
        val.font.size = Pt(12)

    # Candidate summary table
    doc.add_paragraph()  # spacing
    heading = doc.add_paragraph()
    hrun = heading.add_run("Candidates Included")
    hrun.bold = True
    hrun.font.name = "Calibri"
    hrun.font.size = Pt(14)
    hrun.font.color.rgb = NAVY

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    headers = table.rows[0].cells
    for i, text in enumerate(["#", "Candidate", "Match Score"]):
        headers[i].text = text

    for idx, c in enumerate(candidates, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = c.get("name", "Unknown")
        score = c.get("match_score")
        row[2].text = f"{score}%" if score is not None else "N/A"

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Brief page DOCX
# ---------------------------------------------------------------------------

def _build_brief_page(
    candidate_name: str,
    brief_text: str,
    output_path: str,
) -> str:
    """Build a DOCX page from an AI-generated submittal brief."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Candidate name heading
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(4)
    run = name_para.add_run(candidate_name)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    # Subheading
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    srun = sub.add_run("Candidate Brief")
    srun.font.name = "Calibri"
    srun.font.size = Pt(11)
    srun.font.color.rgb = GRAY
    srun.italic = True

    # Brief content — split by lines for formatting
    for line in brief_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        # Detect markdown-style headers
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(line[3:])
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(13)
            r.font.color.rgb = BLUE
        elif line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run(line[2:])
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(14)
            r.font.color.rgb = NAVY
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Calibri"
            r.font.size = Pt(11)

    doc.save(output_path)
    return output_path


def _build_external_brief_page(
    candidate: RecruiterPipelineCandidate,
    output_path: str,
) -> str:
    """Build a brief page for an external (non-platform) candidate."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    name = candidate.external_name or "Unknown Candidate"
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(4)
    run = name_para.add_run(name)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    # Details
    details = [
        ("Current Title", candidate.current_title),
        ("Current Company", candidate.current_company),
        ("Location", candidate.location),
        ("Email", candidate.external_email),
        ("Phone", candidate.external_phone),
        ("LinkedIn", candidate.external_linkedin),
        ("Skills", candidate.skills),
    ]
    for label, value in details:
        if not value:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        lbl = p.add_run(f"{label}: ")
        lbl.bold = True
        lbl.font.name = "Calibri"
        lbl.font.size = Pt(11)
        val = p.add_run(str(value))
        val.font.name = "Calibri"
        val.font.size = Pt(11)

    if candidate.notes:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run("Notes")
        hr.bold = True
        hr.font.name = "Calibri"
        hr.font.size = Pt(13)
        hr.font.color.rgb = BLUE
        np = doc.add_paragraph()
        nr = np.add_run(candidate.notes)
        nr.font.name = "Calibri"
        nr.font.size = Pt(11)

    if candidate.match_score is not None:
        doc.add_paragraph()
        sp = doc.add_paragraph()
        sr = sp.add_run(f"Match Score: {candidate.match_score}%")
        sr.bold = True
        sr.font.name = "Calibri"
        sr.font.size = Pt(12)
        sr.font.color.rgb = BLUE

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Main worker task
# ---------------------------------------------------------------------------

def build_submittal_package_task(package_id: int) -> dict:
    """RQ worker task: build a submittal package PDF.

    Generates AI briefs, assembles DOCX pages, merges into single PDF.
    """
    session = get_session_factory()()
    try:
        pkg = session.get(SubmittalPackage, package_id)
        if not pkg:
            return {"error": "Package not found"}

        job = session.get(RecruiterJob, pkg.recruiter_job_id)
        if not job:
            pkg.status = "failed"
            pkg.error_message = "Recruiter job not found"
            session.commit()
            return {"error": pkg.error_message}

        # Resolve client name
        client_name = None
        if pkg.client_id:
            client = session.get(RecruiterClient, pkg.client_id)
            if client:
                client_name = client.company_name

        candidate_ids = pkg.candidate_ids or []
        pipeline_ids = pkg.pipeline_candidate_ids or []

        # Collect candidate info for cover page
        candidate_summaries: list[dict] = []
        documents: list[dict] = []  # For merge_documents_to_pdf
        tmp_dir = Path(tempfile.mkdtemp(prefix="submittal_"))

        # --- Cover page ---
        # Gather summaries first, then build cover page
        # Platform candidates
        for cid in candidate_ids:
            profile = session.execute(
                select(CandidateProfile).where(CandidateProfile.id == cid)
            ).scalar_one_or_none()
            if not profile:
                continue
            pdata = profile.profile_json or {}
            basics = pdata.get("basics", {})
            name = basics.get("name", f"Candidate #{cid}")

            # Get match score from RecruiterJobCandidate if available
            from app.models.recruiter_job_candidate import RecruiterJobCandidate

            match = session.execute(
                select(RecruiterJobCandidate).where(
                    RecruiterJobCandidate.recruiter_job_id == job.id,
                    RecruiterJobCandidate.candidate_profile_id == cid,
                )
            ).scalar_one_or_none()
            score = round(match.match_score) if match and match.match_score else None

            candidate_summaries.append({
                "name": name,
                "match_score": score,
                "profile_id": cid,
                "type": "platform",
            })

        # External pipeline candidates
        for pid in pipeline_ids:
            pc = session.get(RecruiterPipelineCandidate, pid)
            if not pc:
                continue
            candidate_summaries.append({
                "name": pc.external_name or f"Candidate #{pid}",
                "match_score": round(pc.match_score) if pc.match_score else None,
                "pipeline_id": pid,
                "type": "external",
            })

        if not candidate_summaries:
            pkg.status = "failed"
            pkg.error_message = "No valid candidates found"
            session.commit()
            return {"error": pkg.error_message}

        # Build cover page
        cover_path = str(tmp_dir / "00_cover.docx")
        _build_cover_page(
            job, client_name, pkg.recipient_name,
            candidate_summaries, cover_path,
        )
        documents.append({
            "path": cover_path, "type": "docx", "label": "Cover Page",
        })

        # --- Build per-candidate pages ---
        options = pkg.package_options or {}
        include_briefs = options.get("include_briefs", True)
        include_resumes = options.get("include_resumes", True)

        for idx, cs in enumerate(candidate_summaries, 1):
            prefix = f"{idx:02d}"

            if cs["type"] == "platform":
                cid = cs["profile_id"]

                # AI brief
                if include_briefs:
                    try:
                        from app.services.career_intelligence import (
                            generate_candidate_brief,
                        )

                        brief_result = generate_candidate_brief(
                            candidate_profile_id=cid,
                            employer_job_id=job.id,
                            brief_type="submittal",
                            user_id=None,
                            db=session,
                        )
                        brief_text = brief_result.get("brief_text", "")
                        if brief_text:
                            brief_path = str(
                                tmp_dir / f"{prefix}_brief_{cid}.docx"
                            )
                            _build_brief_page(cs["name"], brief_text, brief_path)
                            documents.append({
                                "path": brief_path,
                                "type": "docx",
                                "label": f"Brief - {cs['name']}",
                            })
                    except Exception:
                        logger.exception(
                            "Failed to generate brief for candidate %s", cid
                        )
                        # Add placeholder
                        brief_path = str(
                            tmp_dir / f"{prefix}_brief_{cid}.docx"
                        )
                        _build_brief_page(
                            cs["name"],
                            "Brief generation was not available for this candidate.",
                            brief_path,
                        )
                        documents.append({
                            "path": brief_path,
                            "type": "docx",
                            "label": f"Brief - {cs['name']}",
                        })

                # Resume
                if include_resumes:
                    resume_path = _find_candidate_resume(
                        session, cid, job.id
                    )
                    if resume_path:
                        suffix = Path(resume_path).suffix.lower()
                        doc_type = "pdf" if suffix == ".pdf" else "docx"
                        documents.append({
                            "path": resume_path,
                            "type": doc_type,
                            "label": f"Resume - {cs['name']}",
                        })

            elif cs["type"] == "external":
                pid = cs["pipeline_id"]
                pc = session.get(RecruiterPipelineCandidate, pid)
                if not pc:
                    continue

                # External brief page
                if include_briefs:
                    brief_path = str(
                        tmp_dir / f"{prefix}_brief_ext_{pid}.docx"
                    )
                    _build_external_brief_page(pc, brief_path)
                    documents.append({
                        "path": brief_path,
                        "type": "docx",
                        "label": f"Brief - {cs['name']}",
                    })

                # External resume
                if include_resumes and pc.external_resume_url:
                    suffix = Path(pc.external_resume_url).suffix.lower()
                    doc_type = "pdf" if suffix == ".pdf" else "docx"
                    documents.append({
                        "path": pc.external_resume_url,
                        "type": doc_type,
                        "label": f"Resume - {cs['name']}",
                    })

        # --- Merge into PDF ---
        from app.services.document_merger import merge_documents_to_pdf

        job_slug = (job.title or "job").replace(" ", "_")[:30]
        client_slug = (
            client_name or job.client_company_name or "client"
        ).replace(" ", "_")[:20]
        filename = f"Submittal_{client_slug}_{job_slug}_{pkg.id}"

        merged_url = merge_documents_to_pdf(documents, filename)

        # --- Create CandidateSubmission records for platform candidates ---
        for cs in candidate_summaries:
            if cs["type"] != "platform":
                continue
            cid = cs["profile_id"]
            # Check if submission already exists
            existing = session.execute(
                select(CandidateSubmission).where(
                    CandidateSubmission.recruiter_job_id == job.id,
                    CandidateSubmission.candidate_profile_id == cid,
                )
            ).scalar_one_or_none()
            if existing:
                # Link to this package if not already linked
                if not existing.submittal_package_id:
                    existing.submittal_package_id = pkg.id
                continue

            sub = CandidateSubmission(
                recruiter_job_id=job.id,
                candidate_profile_id=cid,
                recruiter_profile_id=pkg.recruiter_profile_id,
                submittal_package_id=pkg.id,
                status="submitted",
            )
            session.add(sub)

        # Update package
        pkg.merged_pdf_url = merged_url
        pkg.status = "ready"
        session.commit()

        # Clean up temp files
        import shutil

        shutil.rmtree(tmp_dir, ignore_errors=True)

        return {
            "package_id": pkg.id,
            "status": "ready",
            "merged_pdf_url": merged_url,
            "candidates": len(candidate_summaries),
        }

    except Exception as e:
        logger.exception("Failed to build submittal package %s", package_id)
        try:
            session.rollback()
            pkg = session.get(SubmittalPackage, package_id)
            if pkg:
                pkg.status = "failed"
                pkg.error_message = str(e)[:1000]
                session.commit()
        except Exception:
            logger.exception("Failed to update package status")
        return {"error": str(e)}
    finally:
        session.close()


# ---------------------------------------------------------------------------
# Email sending
# ---------------------------------------------------------------------------

def send_submittal_email(session: Session, package_id: int) -> dict:
    """Send the submittal package PDF to the client via email."""
    pkg = session.get(SubmittalPackage, package_id)
    if not pkg:
        raise ValueError("Package not found")
    if pkg.status not in ("ready", "sent"):
        raise ValueError(f"Package is not ready (status={pkg.status})")
    if not pkg.merged_pdf_url:
        raise ValueError("No merged PDF available")

    job = session.get(RecruiterJob, pkg.recruiter_job_id)
    job_title = job.title if job else "Position"

    subject = pkg.cover_email_subject or (
        f"Candidate Submittal — {job_title}"
    )
    body = pkg.cover_email_body or (
        f"<p>Please find attached our candidate submittal package for "
        f"<strong>{job_title}</strong>.</p>"
        f"<p>This package includes candidate briefs and supporting "
        f"documentation for your review.</p>"
        f"<p>Please don't hesitate to reach out with any questions.</p>"
    )

    # Download PDF and send
    from app.services.storage import download_to_tempfile

    local_pdf = download_to_tempfile(pkg.merged_pdf_url, suffix=".pdf")
    pdf_bytes = local_pdf.read_bytes()
    pdf_filename = Path(pkg.merged_pdf_url).stem + ".pdf"

    from app.services.email import send_submittal_package_email

    message_id = send_submittal_package_email(
        to_email=pkg.recipient_email,
        to_name=pkg.recipient_name,
        subject=subject,
        body_html=body,
        pdf_bytes=pdf_bytes,
        pdf_filename=pdf_filename,
    )

    pkg.status = "sent"
    pkg.sent_at = datetime.now(UTC)
    pkg.email_message_id = message_id or ""
    pkg.cover_email_subject = subject
    pkg.cover_email_body = body

    # Update job status to submitted
    if job and job.status == "active":
        job.status = "submitted"

    session.commit()

    # Clean up temp file
    try:
        local_pdf.unlink(missing_ok=True)
    except Exception:
        pass

    return {
        "status": "sent",
        "email_message_id": message_id,
    }


# ---------------------------------------------------------------------------
# List packages
# ---------------------------------------------------------------------------

def list_packages(
    session: Session, recruiter_profile_id: int, recruiter_job_id: int
) -> list[SubmittalPackage]:
    """List submittal packages for a recruiter job."""
    return list(
        session.execute(
            select(SubmittalPackage)
            .where(
                SubmittalPackage.recruiter_profile_id == recruiter_profile_id,
                SubmittalPackage.recruiter_job_id == recruiter_job_id,
            )
            .order_by(SubmittalPackage.created_at.desc())
        )
        .scalars()
        .all()
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_candidate_resume(
    session: Session, candidate_profile_id: int, recruiter_job_id: int
) -> str | None:
    """Find the best available resume for a candidate.

    Priority: tailored resume for this job > original resume document.
    """
    profile = session.get(CandidateProfile, candidate_profile_id)
    if not profile:
        return None

    # Check for tailored resume (linked via user_id and a matching job)
    if profile.user_id:
        # Look for any tailored resume for this user
        tailored = session.execute(
            select(TailoredResume)
            .where(TailoredResume.user_id == profile.user_id)
            .order_by(TailoredResume.created_at.desc())
            .limit(1)
        ).scalar_one_or_none()
        if tailored and tailored.docx_url:
            return tailored.docx_url

    # Fall back to original resume document
    if profile.resume_document_id:
        doc = session.get(ResumeDocument, profile.resume_document_id)
        if doc and doc.path and not doc.deleted_at:
            return doc.path

    return None

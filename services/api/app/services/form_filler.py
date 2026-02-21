"""Form Filler — populates employer DOCX forms from candidate profiles.

Works on a COPY of the original template. Preserves all formatting
(fonts, borders, shading, alignment). Never fills signature, rate,
or vendor fields.
"""

import logging
import re
import shutil

from docx import Document

logger = logging.getLogger(__name__)


def fill_form(
    template_path: str,
    parsed_structure: dict,
    profile_json: dict,
    output_path: str,
) -> dict:
    """Fill a DOCX form from profile data.

    Args:
        template_path: Path to the original DOCX template.
        parsed_structure: Output of form_parser.parse_form_document().
        profile_json: The candidate's profile_json dict.
        output_path: Where to write the filled DOCX.

    Returns:
        {
            "filled_count": int,
            "unfilled_fields": [{"label": str, "reason": str}],
            "gaps_detected": [{"label": str, "description": str}],
            "output_path": str,
        }
    """
    # Always work on a copy
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    filled = 0
    unfilled: list[dict] = []
    gaps: list[dict] = []

    sections = parsed_structure.get("sections", [])
    for section in sections:
        section_type = section.get("type", "other")

        if section_type == "skills_matrix":
            result = _fill_skills_matrix(doc, section, profile_json)
        elif section_type == "reference_form":
            refs = profile_json.get("references", [])
            active_refs = [r for r in refs if r.get("is_active", True)]
            result = _fill_reference_form(doc, section, active_refs)
        elif section_type in ("acknowledgement", "right_to_represent"):
            result = _fill_acknowledgement(doc, section, profile_json)
        else:
            result = _fill_generic_section(doc, section, profile_json)

        filled += result.get("filled", 0)
        unfilled.extend(result.get("unfilled", []))
        gaps.extend(result.get("gaps", []))

    doc.save(output_path)

    return {
        "filled_count": filled,
        "unfilled_fields": unfilled,
        "gaps_detected": gaps,
        "output_path": output_path,
    }


def fill_cell_preserving_format(cell, value: str) -> None:
    """Write value into a table cell while preserving existing formatting."""
    if not cell.paragraphs:
        return

    para = cell.paragraphs[0]

    if para.runs:
        # Preserve the first run's formatting
        run = para.runs[0]
        run.text = value
        # Clear any additional runs
        for extra_run in para.runs[1:]:
            extra_run.text = ""
    else:
        # No runs — add one copying paragraph-level formatting
        run = para.add_run(value)

    # Preserve existing cell shading / borders (they live on the cell XML)
    # No need to modify — we only touched the run text.


def _fill_skills_matrix(doc: Document, section: dict, profile_json: dict) -> dict:
    """Fill a qualification/skills matrix table."""
    filled = 0
    unfilled: list[dict] = []
    gaps: list[dict] = []

    table_idx = section.get("table_index")
    if table_idx is None or table_idx >= len(doc.tables):
        return {"filled": 0, "unfilled": [], "gaps": []}

    table = doc.tables[table_idx]
    skill_years = profile_json.get("skill_years", {})
    profile_skills = [s.lower() for s in profile_json.get("skills", [])]

    for field in section.get("fields", []):
        if field.get("never_fill"):
            unfilled.append({"label": field["label"], "reason": "never_fill"})
            continue

        skill_label = field["label"]
        row_idx = field.get("row", 0)
        years_col = field.get("years_col", 1)

        if row_idx >= len(table.rows):
            continue

        match = _match_skill_to_profile(skill_label, skill_years, profile_skills)
        if match:
            years_val = str(match["years_experience"])
            cell = table.rows[row_idx].cells[years_col]
            fill_cell_preserving_format(cell, years_val)
            filled += 1
        else:
            gaps.append(
                {
                    "label": skill_label,
                    "description": f"No matching skill/years found for '{skill_label}'",
                }
            )
            unfilled.append({"label": skill_label, "reason": "no_match"})

    return {"filled": filled, "unfilled": unfilled, "gaps": gaps}


def _fill_reference_form(doc: Document, section: dict, references: list[dict]) -> dict:
    """Fill reference fields from the candidate's references list."""
    filled = 0
    unfilled: list[dict] = []
    gaps: list[dict] = []

    table_idx = section.get("table_index")
    if table_idx is None or table_idx >= len(doc.tables):
        return {"filled": 0, "unfilled": [], "gaps": []}

    table = doc.tables[table_idx]

    for field in section.get("fields", []):
        if field.get("never_fill"):
            unfilled.append({"label": field["label"], "reason": "never_fill"})
            continue

        ref_idx = field.get("ref_index", 0)
        ref_field = field.get("ref_field", "")
        row_idx = field.get("row", 0)
        col_idx = field.get("col", 1)

        if ref_idx >= len(references):
            gaps.append(
                {
                    "label": field["label"],
                    "description": f"Reference #{ref_idx + 1} not provided",
                }
            )
            unfilled.append({"label": field["label"], "reason": "missing_reference"})
            continue

        ref = references[ref_idx]
        value = str(ref.get(ref_field, ""))

        row_ok = row_idx < len(table.rows)
        col_ok = row_ok and col_idx < len(table.rows[row_idx].cells)
        if value and col_ok:
            cell = table.rows[row_idx].cells[col_idx]
            fill_cell_preserving_format(cell, value)
            filled += 1
        else:
            unfilled.append({"label": field["label"], "reason": "empty_value"})

    return {"filled": filled, "unfilled": unfilled, "gaps": gaps}


def _fill_acknowledgement(doc: Document, section: dict, profile_json: dict) -> dict:
    """Fill acknowledgement fields — only candidate name, never signature/date."""
    filled = 0
    unfilled: list[dict] = []

    basics = profile_json.get("basics", {})
    name = basics.get("name", "")

    for field in section.get("fields", []):
        if field.get("never_fill"):
            unfilled.append({"label": field["label"], "reason": "never_fill"})
            continue

        profile_path = field.get("profile_path")
        if profile_path and profile_path == ["basics", "name"] and name:
            para_idx = field.get("paragraph_index")
            if para_idx is not None and para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                # Replace underscores with name
                new_text = re.sub(r"_+", name, para.text)
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""
                filled += 1
            continue

        unfilled.append({"label": field.get("label", ""), "reason": "not_fillable"})

    return {"filled": filled, "unfilled": unfilled, "gaps": []}


def _fill_generic_section(doc: Document, section: dict, profile_json: dict) -> dict:
    """Fill generic label-value fields from profile."""
    filled = 0
    unfilled: list[dict] = []

    table_idx = section.get("table_index")
    if table_idx is None or table_idx >= len(doc.tables):
        return {"filled": 0, "unfilled": [], "gaps": []}

    table = doc.tables[table_idx]

    for field in section.get("fields", []):
        if field.get("never_fill"):
            unfilled.append({"label": field["label"], "reason": "never_fill"})
            continue

        profile_path = field.get("profile_path")
        if not profile_path:
            unfilled.append({"label": field["label"], "reason": "unmapped"})
            continue

        # Resolve value from profile
        value = _resolve_profile_value(profile_json, profile_path)
        if not value:
            unfilled.append({"label": field["label"], "reason": "empty"})
            continue

        row_idx = field.get("row", 0)
        col_idx = field.get("col", 1)
        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
            cell = table.rows[row_idx].cells[col_idx]
            fill_cell_preserving_format(cell, str(value))
            filled += 1

    return {"filled": filled, "unfilled": unfilled, "gaps": []}


def _resolve_profile_value(profile_json: dict, path: list[str]) -> str | None:
    """Walk a profile_json path to retrieve a value."""
    current = profile_json
    for key in path:
        if isinstance(current, dict):
            current = current.get(key)
        elif isinstance(current, list):
            try:
                idx = int(key)
                current = current[idx] if idx < len(current) else None
            except (ValueError, IndexError):
                return None
        else:
            return None
        if current is None:
            return None
    return str(current) if current else None


def _match_skill_to_profile(
    description: str,
    skill_years: dict,
    profile_skills: list[str],
) -> dict | None:
    """Match a form skill description to profile skill_years.

    Uses exact match first, then keyword substring matching.
    """
    desc_lower = description.lower().strip()

    # Exact match in skill_years
    for skill_name, data in skill_years.items():
        if skill_name.lower() == desc_lower:
            return data

    # Substring match
    for skill_name, data in skill_years.items():
        sn = skill_name.lower()
        if sn in desc_lower or desc_lower in sn:
            return data

    # Check if skill exists in profile (even without years data)
    for ps in profile_skills:
        if ps in desc_lower or desc_lower in ps:
            return {"years_experience": 1, "years_experience_source": "inferred"}

    return None

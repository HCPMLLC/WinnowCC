from __future__ import annotations

import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.candidate_profile import CandidateProfile
from app.models.job import Job
from app.models.match import Match
from app.models.tailored_resume import TailoredResume
from app.services.cover_letter_scoring import compute_cover_letter_score
from app.services.matching import recalculate_interview_probability
from app.services.storage import upload_file


def create_tailored_docs(
    session: Session, user_id: int, job_id: int, profile_version: int
) -> TailoredResume:
    job = session.get(Job, job_id)
    if job is None:
        raise ValueError("Job not found.")

    profile = _get_profile(session, user_id, profile_version)
    if profile is None:
        raise ValueError("Candidate profile not found.")

    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    resume_fname = f"resume_{user_id}_{job_id}_{timestamp}.docx"
    cover_fname = f"cover_{user_id}_{job_id}_{timestamp}.docx"

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        resume_local = tmp / resume_fname
        cover_local = tmp / cover_fname

        _build_resume_doc(resume_local, job, profile.profile_json)
        _build_cover_letter_doc(cover_local, job, profile.profile_json)

        # Compute cover letter score on local temp file before upload
        cover_letter_score = compute_cover_letter_score(
            cover_letter_path=cover_local,
            job_description=job.description_text,
            company_name=job.company,
            hiring_manager_name=job.hiring_manager_name,
        )

        # Upload to GCS or copy to local data dir
        resume_stored = upload_file(resume_local, "tailored/", resume_fname)
        cover_stored = upload_file(cover_local, "tailored/", cover_fname)

    # Update match with cover letter score if exists
    match = session.execute(
        select(Match).where(
            Match.user_id == user_id,
            Match.job_id == job_id,
            Match.profile_version == profile.version,
        )
    ).scalar_one_or_none()
    if match:
        match.cover_letter_score = cover_letter_score
        match.interview_probability = recalculate_interview_probability(match)

    tailored = TailoredResume(
        user_id=user_id,
        job_id=job_id,
        profile_version=profile.version,
        docx_url=resume_stored,
        cover_letter_url=cover_stored,
        change_log={
            "job_title": job.title,
            "matched_skills": profile.profile_json.get("skills", []),
            "cover_letter_score": cover_letter_score,
        },
    )
    session.add(tailored)
    session.commit()
    session.refresh(tailored)
    return tailored


def _get_profile(
    session: Session, user_id: int, profile_version: int
) -> CandidateProfile | None:
    if profile_version > 0:
        stmt = select(CandidateProfile).where(
            CandidateProfile.user_id == user_id,
            CandidateProfile.version == profile_version,
        )
    else:
        stmt = (
            select(CandidateProfile)
            .where(CandidateProfile.user_id == user_id)
            .order_by(CandidateProfile.version.desc())
            .limit(1)
        )
    return session.execute(stmt).scalars().first()


def _build_resume_doc(path: Path, job: Job, profile_json: dict) -> None:
    doc = Document()
    _set_default_font(doc)

    _add_heading(doc, job.title, level=0)

    basics = profile_json.get("basics", {}) if isinstance(profile_json, dict) else {}
    contact = " | ".join(
        [value for value in [basics.get("email"), basics.get("phone"), basics.get("location")] if value]
    )
    if contact:
        doc.add_paragraph(contact)

    _add_section_heading(doc, "Work Experience")
    for item in profile_json.get("experience", []) or []:
        title = item.get("title") or "Role"
        company = item.get("company") or ""
        header = f"{title} - {company}".strip(" -")
        doc.add_paragraph(header)
        dates = " - ".join(
            [value for value in [item.get("start_date"), item.get("end_date")] if value]
        )
        if dates:
            doc.add_paragraph(dates)
        bullets = item.get("bullets") or []
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    _add_section_heading(doc, "Education")
    for item in profile_json.get("education", []) or []:
        school = item.get("school") or "School"
        degree = item.get("degree") or ""
        field = item.get("field") or ""
        detail = " ".join([degree, field]).strip()
        line = f"{school} - {detail}".strip(" -")
        doc.add_paragraph(line)

    _add_section_heading(doc, "Certifications")
    doc.add_paragraph("None listed.")

    _add_section_heading(doc, "Skills")
    skills = [s for s in (profile_json.get("skills", []) or []) if isinstance(s, str)]
    matched = [s for s in skills if s.lower() in job.description_text.lower()]
    remaining = [s for s in skills if s not in matched]
    ordered = matched + remaining
    doc.add_paragraph(", ".join(ordered))

    doc.save(str(path))


def _build_cover_letter_doc(path: Path, job: Job, profile_json: dict) -> None:
    doc = Document()
    _set_default_font(doc)

    name = (profile_json.get("basics", {}) or {}).get("name") or "Candidate"
    hiring_manager = job.hiring_manager_name or "Hiring Manager"
    doc.add_paragraph(f"Dear {hiring_manager},")

    doc.add_paragraph(
        f"I am excited to apply for the {job.title} role at {job.company}. "
        "My background aligns closely with your needs, and I would welcome the chance "
        "to contribute immediately."
    )

    doc.add_paragraph(
        f"I am especially interested in {job.company}'s focus on [company mission or recent initiative]."
    )

    requirements = _top_requirements(job.description_text)
    if requirements:
        doc.add_paragraph("Key alignments:")
        for req in requirements[:3]:
            doc.add_paragraph(req, style="List Bullet")

    doc.add_paragraph(
        "I appreciate the opportunity to bring my experience to your team. "
        "If helpful, I can share additional examples of impact and walk through "
        "how I would approach your immediate priorities."
    )

    doc.add_paragraph(f"Sincerely,\n{name}")
    doc.save(str(path))


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.bold = True


def _add_section_heading(doc: Document, title: str) -> None:
    _add_heading(doc, title, level=1)


def _top_requirements(description: str) -> list[str]:
    sentences = [s.strip() for s in description.split(".") if s.strip()]
    if not sentences:
        return []
    return sentences[:3]

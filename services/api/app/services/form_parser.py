"""Form Parser — extracts fillable structure from employer DOCX forms.

Detects qualification tables, reference forms, acknowledgement sections,
and maps form fields to candidate profile_json paths.
"""

import logging
import re
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

# Section type constants
SKILLS_MATRIX = "skills_matrix"
REFERENCE_FORM = "reference_form"
ACKNOWLEDGEMENT = "acknowledgement"
RIGHT_TO_REPRESENT = "right_to_represent"
OTHER = "other"

# Fields that must NEVER be auto-filled
NEVER_FILL_PATTERNS = [
    re.compile(r"\bsignature\b", re.IGNORECASE),
    re.compile(r"\bsign\s*here\b", re.IGNORECASE),
    re.compile(r"\bdate\b.*\bsignature\b", re.IGNORECASE),
    re.compile(r"\bsignature\b.*\bdate\b", re.IGNORECASE),
    re.compile(r"\bhourly\s*rate\b", re.IGNORECASE),
    re.compile(r"\bbill\s*rate\b", re.IGNORECASE),
    re.compile(r"\bvendor\b", re.IGNORECASE),
    re.compile(r"\bsubcontractor\b", re.IGNORECASE),
    re.compile(r"\bprime\s*contractor\b", re.IGNORECASE),
]

# Profile field mapping: form label keywords -> profile_json path
_FIELD_MAP: dict[str, list[str]] = {
    # Basics
    "candidate name": ["basics", "name"],
    "full name": ["basics", "name"],
    "first name": ["basics", "first_name"],
    "last name": ["basics", "last_name"],
    "name": ["basics", "name"],
    "email": ["basics", "email"],
    "phone": ["basics", "phone"],
    "location": ["basics", "location"],
    # Reference fields
    "reference name": ["references", "*", "name"],
    "reference title": ["references", "*", "title"],
    "reference company": ["references", "*", "company"],
    "reference phone": ["references", "*", "phone"],
    "reference email": ["references", "*", "email"],
    "relationship": ["references", "*", "relationship"],
    "years known": ["references", "*", "years_known"],
}

# Skills-matrix header patterns
_SKILL_HEADER_RE = re.compile(
    r"(skill|technology|tool|competenc|qualification|software|platform|"
    r"certification|experience area)",
    re.IGNORECASE,
)
_YEARS_HEADER_RE = re.compile(
    r"(year|yr|yrs|experience|months|duration|how long)", re.IGNORECASE
)

# Reference section patterns
_REF_SECTION_RE = re.compile(r"(reference|professional\s+reference)", re.IGNORECASE)

# Acknowledgement patterns
_ACK_SECTION_RE = re.compile(
    r"(acknowledge|acknowledgement|agreement|consent|right\s+to\s+represent|"
    r"authorization|attest)",
    re.IGNORECASE,
)


def parse_form_document(file_path: str, job_id: int) -> dict:
    """Parse a DOCX form and extract its fillable structure.

    Returns:
        {
            "job_id": int,
            "sections": [{"type": str, "fields": [...], "table_index": int}],
            "total_fields": int,
            "auto_fillable": int,
            "needs_manual": int,
        }
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Form not found: {file_path}")

    doc = Document(str(path))
    sections = _detect_form_sections(doc)

    total = 0
    fillable = 0
    manual = 0
    for section in sections:
        for field in section.get("fields", []):
            total += 1
            if field.get("never_fill"):
                manual += 1
            elif field.get("profile_path"):
                fillable += 1
            else:
                manual += 1

    return {
        "job_id": job_id,
        "sections": sections,
        "total_fields": total,
        "auto_fillable": fillable,
        "needs_manual": manual,
    }


def _detect_form_sections(doc: Document) -> list[dict]:
    """Scan document tables and paragraphs for fillable sections."""
    sections: list[dict] = []

    # Scan tables
    for idx, table in enumerate(doc.tables):
        section_type = _classify_table(table)
        fields = _extract_table_fields(table, section_type)
        if fields:
            sections.append(
                {"type": section_type, "table_index": idx, "fields": fields}
            )

    # Scan paragraphs for acknowledgement sections
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if _ACK_SECTION_RE.search(text):
            section_type = (
                RIGHT_TO_REPRESENT
                if re.search(r"right\s+to\s+represent", text, re.IGNORECASE)
                else ACKNOWLEDGEMENT
            )
            fields = _extract_paragraph_fields(doc.paragraphs, idx)
            if fields:
                sections.append(
                    {"type": section_type, "paragraph_index": idx, "fields": fields}
                )

    return sections


def _classify_table(table) -> str:
    """Classify a table as skills_matrix, reference_form, or other."""
    header_text = ""
    if table.rows:
        header_text = " ".join(cell.text for cell in table.rows[0].cells)

    if _SKILL_HEADER_RE.search(header_text) and _YEARS_HEADER_RE.search(header_text):
        return SKILLS_MATRIX
    if _REF_SECTION_RE.search(header_text):
        return REFERENCE_FORM

    # Check second column for years indicators
    if len(table.columns) >= 2:
        col_text = " ".join(
            table.rows[r].cells[1].text
            for r in range(min(3, len(table.rows)))
            if len(table.rows[r].cells) > 1
        )
        if _YEARS_HEADER_RE.search(col_text):
            return SKILLS_MATRIX

    return OTHER


def _extract_table_fields(table, section_type: str) -> list[dict]:
    """Extract fillable fields from a table."""
    fields: list[dict] = []

    if section_type == SKILLS_MATRIX:
        return _extract_skills_matrix_fields(table)
    if section_type == REFERENCE_FORM:
        return _extract_reference_fields(table)

    # Generic table: label-value pairs
    for row_idx, row in enumerate(table.rows):
        cells = row.cells
        for cell_idx in range(0, len(cells) - 1, 2):
            label = cells[cell_idx].text.strip()
            value = (
                cells[cell_idx + 1].text.strip() if cell_idx + 1 < len(cells) else ""
            )
            if label and len(label) < 100:
                never_fill = _is_never_fill(label)
                profile_path = (
                    None if never_fill else _map_field_to_profile(label, section_type)
                )
                fields.append(
                    {
                        "label": label,
                        "current_value": value,
                        "row": row_idx,
                        "col": cell_idx + 1,
                        "profile_path": profile_path,
                        "never_fill": never_fill,
                    }
                )
    return fields


def _extract_skills_matrix_fields(table) -> list[dict]:
    """Extract skill rows from a qualification/skills matrix table."""
    fields: list[dict] = []
    if not table.rows:
        return fields

    # Find skill and years columns from header
    header_cells = table.rows[0].cells
    skill_col = None
    years_col = None
    for i, cell in enumerate(header_cells):
        text = cell.text.strip()
        if _SKILL_HEADER_RE.search(text) and skill_col is None:
            skill_col = i
        elif _YEARS_HEADER_RE.search(text) and years_col is None:
            years_col = i

    if skill_col is None:
        skill_col = 0
    if years_col is None:
        years_col = 1 if len(header_cells) > 1 else 0

    for row_idx in range(1, len(table.rows)):
        cells = table.rows[row_idx].cells
        if skill_col < len(cells):
            skill_text = cells[skill_col].text.strip()
            years_text = cells[years_col].text.strip() if years_col < len(cells) else ""
            if skill_text:
                fields.append(
                    {
                        "label": skill_text,
                        "field_type": "skill_years",
                        "current_value": years_text,
                        "row": row_idx,
                        "skill_col": skill_col,
                        "years_col": years_col,
                        "profile_path": ["skill_years", skill_text],
                        "never_fill": False,
                    }
                )
    return fields


def _extract_reference_fields(table) -> list[dict]:
    """Extract reference fields from a reference form table."""
    fields: list[dict] = []
    ref_index = 0

    for row_idx, row in enumerate(table.rows):
        cells = row.cells
        for cell_idx, cell in enumerate(cells):
            text = cell.text.strip().rstrip(":").lower()
            if not text or len(text) > 80:
                continue

            # Detect reference field labels
            ref_kws = ("name", "title", "company", "phone", "email", "relationship")
            if any(kw in text for kw in ref_kws):
                value_col = cell_idx + 1
                current = (
                    cells[value_col].text.strip() if value_col < len(cells) else ""
                )

                field_key = text.split()[-1] if text.split() else text
                never_fill = _is_never_fill(text)

                fields.append(
                    {
                        "label": text,
                        "field_type": "reference",
                        "current_value": current,
                        "row": row_idx,
                        "col": value_col,
                        "ref_index": ref_index,
                        "ref_field": field_key,
                        "profile_path": ["references", str(ref_index), field_key],
                        "never_fill": never_fill,
                    }
                )

            # New reference block detection (when we see "name" again)
            if "name" in text and row_idx > 0:
                prev_had_name = any(
                    f.get("ref_field") == "name" and f.get("ref_index") == ref_index
                    for f in fields[:-1]
                )
                if prev_had_name:
                    ref_index += 1
                    # Update the last added field's ref_index
                    if fields:
                        fields[-1]["ref_index"] = ref_index
                        fields[-1]["profile_path"] = [
                            "references",
                            str(ref_index),
                            fields[-1].get("ref_field", ""),
                        ]

    return fields


def _extract_paragraph_fields(paragraphs: list, start_idx: int) -> list[dict]:
    """Extract fillable fields from acknowledgement paragraph sections."""
    fields: list[dict] = []

    for i in range(start_idx, min(start_idx + 10, len(paragraphs))):
        text = paragraphs[i].text.strip()
        if not text:
            continue

        # Look for blank lines or underscores indicating fillable fields
        if "____" in text or "______" in text:
            label = re.sub(r"_+", "", text).strip().rstrip(":")
            never_fill = _is_never_fill(label)
            profile_path = (
                None if never_fill else _map_field_to_profile(label, ACKNOWLEDGEMENT)
            )
            fields.append(
                {
                    "label": label or "field",
                    "paragraph_index": i,
                    "profile_path": profile_path,
                    "never_fill": never_fill,
                }
            )

    return fields


def _is_never_fill(label: str) -> bool:
    """Check if a field should never be auto-filled."""
    return any(p.search(label) for p in NEVER_FILL_PATTERNS)


def _map_field_to_profile(field_label: str, section_type: str) -> list[str] | None:
    """Map a form field label to a profile_json path."""
    label_lower = field_label.lower().strip().rstrip(":")

    # Direct mapping lookup
    for pattern, path in _FIELD_MAP.items():
        if pattern in label_lower:
            return path

    # Skills matrix fields map to skill_years
    if section_type == SKILLS_MATRIX:
        return ["skill_years", field_label]

    return None

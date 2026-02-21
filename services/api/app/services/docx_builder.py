"""ATS-safe DOCX builder for tailored resumes and cover letters.

Builds Word documents from structured JSON content produced by the LLM
tailoring pipeline. All formatting follows ATS-safe rules:
- Single column layout, no tables for layout
- Standard fonts (Calibri), standard heading styles
- Real bullet list formatting (not unicode characters)
- No text boxes, shapes, images, or graphics
- Name and contact info in the body (not header/footer)
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


def build_tailored_docx(tailored_content: dict, output_path: str) -> str:
    """Build an ATS-safe DOCX from the tailored content JSON.

    Args:
        tailored_content: Validated JSON from LLM with resume sections.
            Expected keys: basics, professional_summary, experience,
            education, skills, certifications
        output_path: Where to save the DOCX file.

    Returns:
        The output_path for storage.
    """
    doc = Document()
    _set_default_font(doc)
    _set_margins(doc)

    basics = tailored_content.get("basics", {})

    # --- Candidate Name (centered, bold, 16pt) ---
    name = basics.get("name") or "Candidate"
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.name = "Calibri"
    name_run.font.size = Pt(16)

    # --- Contact info (centered, 10pt) ---
    contact_parts = [
        basics.get("email"),
        basics.get("phone"),
        basics.get("location"),
        basics.get("linkedin"),
    ]
    contact_line = " | ".join(p for p in contact_parts if p)
    if contact_line:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_before = Pt(0)
        contact_para.paragraph_format.space_after = Pt(6)
        contact_run = contact_para.add_run(contact_line)
        contact_run.font.name = "Calibri"
        contact_run.font.size = Pt(10)

    # --- Professional Summary ---
    summary = tailored_content.get("professional_summary")
    if summary:
        _add_section_heading(doc, "Professional Summary")
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.name = "Calibri"
        summary_run.font.size = Pt(11)

    # --- Experience ---
    experience = tailored_content.get("experience", [])
    if experience:
        _add_section_heading(doc, "Experience")
        for exp in experience:
            company = exp.get("company") or ""
            title = exp.get("title") or ""
            location = exp.get("location") or ""
            start_date = exp.get("start_date") or ""
            end_date = exp.get("end_date") or ""

            # Company name — bold
            if company:
                comp_para = doc.add_paragraph()
                comp_para.paragraph_format.space_after = Pt(0)
                comp_para.paragraph_format.space_before = Pt(6)
                comp_text = company
                if location:
                    comp_text += f" — {location}"
                comp_run = comp_para.add_run(comp_text)
                comp_run.bold = True
                comp_run.font.name = "Calibri"
                comp_run.font.size = Pt(11)

            # Title — italic | dates
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(2)
            title_run = title_para.add_run(title)
            title_run.italic = True
            title_run.font.name = "Calibri"
            title_run.font.size = Pt(11)
            dates_str = ""
            if start_date and end_date:
                dates_str = f" | {start_date} – {end_date}"
            elif start_date:
                dates_str = f" | {start_date} – Present"
            if dates_str:
                dates_run = title_para.add_run(dates_str)
                dates_run.font.name = "Calibri"
                dates_run.font.size = Pt(11)

            # Bullets
            for bullet in exp.get("bullets", []):
                if bullet and bullet.strip():
                    bp = doc.add_paragraph(bullet.strip(), style="List Bullet")
                    for run in bp.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(11)

    # --- Education ---
    education = tailored_content.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            degree = edu.get("degree") or ""
            school = edu.get("school") or ""
            field = edu.get("field") or ""
            graduation_year = edu.get("graduation_year") or edu.get("end_date") or ""

            parts = []
            if degree:
                parts.append(degree)
            if field:
                parts.append(f"in {field}")
            degree_line = " ".join(parts)

            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_after = Pt(2)
            if degree_line:
                deg_run = edu_para.add_run(degree_line)
                deg_run.bold = True
                deg_run.font.name = "Calibri"
                deg_run.font.size = Pt(11)
            if school:
                if degree_line:
                    edu_para.add_run(" | ").font.size = Pt(11)
                school_run = edu_para.add_run(school)
                school_run.font.name = "Calibri"
                school_run.font.size = Pt(11)
            if graduation_year:
                edu_para.add_run(f" | {graduation_year}").font.size = Pt(11)

    # --- Skills ---
    skills = tailored_content.get("skills", {})
    if skills:
        _add_section_heading(doc, "Skills")
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                label = category.replace("_", " ").title()
                skills_para = doc.add_paragraph()
                skills_para.paragraph_format.space_after = Pt(2)
                label_run = skills_para.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.name = "Calibri"
                label_run.font.size = Pt(11)
                skills_text = ", ".join(skill_list)
                skills_run = skills_para.add_run(skills_text)
                skills_run.font.name = "Calibri"
                skills_run.font.size = Pt(11)
        elif isinstance(skills, list):
            skills_para = doc.add_paragraph()
            skills_run = skills_para.add_run(", ".join(skills))
            skills_run.font.name = "Calibri"
            skills_run.font.size = Pt(11)

    # --- Certifications ---
    certifications = tailored_content.get("certifications", [])
    if certifications:
        _add_section_heading(doc, "Certifications")
        for cert in certifications:
            if isinstance(cert, dict):
                cert_name = cert.get("name") or ""
                cert_year = cert.get("year") or cert.get("date_obtained") or ""
                cert_line = cert_name
                if cert_year:
                    cert_line += f" — {cert_year}"
                if cert_line:
                    doc.add_paragraph(cert_line, style="List Bullet")
            elif isinstance(cert, str) and cert:
                doc.add_paragraph(cert, style="List Bullet")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def build_cover_letter_docx(
    cover_letter_content: dict,
    candidate_basics: dict,
    job_info: dict,
    output_path: str,
) -> str:
    """Build a cover letter DOCX from structured JSON content.

    Formatting follows the PROMPT17 spec (Reformatted_Cover_Letter.docx):
    - Body: Cambria 11pt, line spacing 276 twips (~1.15), after=200 default
    - Heading: Calibri 14pt bold #365F91
    - Margins: 1" top/bottom, 1.25" left/right
    - Contact separator: bullet \u2022
    - Blank paragraphs for section gaps (inherit default after=200)
    - Sign-off: two separate paragraphs

    Args:
        cover_letter_content: JSON with greeting, body_paragraphs, closing, sign_off.
        candidate_basics: Candidate basics (name, email, phone, location, linkedin).
        job_info: Job details (company, title, hiring_manager, address).
        output_path: Where to save the DOCX file.

    Returns:
        The output_path for storage.
    """
    doc = Document()
    _set_cover_letter_defaults(doc)
    _set_cover_letter_margins(doc)

    name = candidate_basics.get("name") or "Candidate"

    # 1. Heading: candidate name
    _add_cover_letter_heading1(doc, name)

    # 2. Contact line: City, ST \u2022 Phone \u2022 Email \u2022 LinkedIn
    contact_parts = []
    location = candidate_basics.get("location")
    if location:
        contact_parts.append(location)
    phone = candidate_basics.get("phone")
    if phone:
        contact_parts.append(phone)
    email = candidate_basics.get("email")
    if email:
        contact_parts.append(email)
    linkedin = candidate_basics.get("linkedin")
    if linkedin:
        contact_parts.append(linkedin)
    if contact_parts:
        _add_cl_body_para(doc, " \u2022 ".join(contact_parts))

    # 3. Blank separator
    _add_cl_blank_separator(doc)

    # 4. Date
    _add_cl_body_para(doc, datetime.now().strftime("%B %d, %Y"))

    # 5. Blank separator
    _add_cl_blank_separator(doc)

    # 6-9. Recipient address block
    hiring_manager = job_info.get("hiring_manager")
    if hiring_manager:
        _add_cl_body_para(doc, hiring_manager)
    company = job_info.get("company") or ""
    if company:
        _add_cl_body_para(doc, company)
    address = job_info.get("address")
    if address:
        if isinstance(address, list):
            for line in address:
                if line and line.strip():
                    _add_cl_body_para(doc, line.strip())
        elif isinstance(address, str) and address.strip():
            for line in address.strip().splitlines():
                if line.strip():
                    _add_cl_body_para(doc, line.strip())

    # 10. Blank separator
    _add_cl_blank_separator(doc)

    # 11. Greeting
    greeting = cover_letter_content.get("greeting") or "Dear Hiring Manager,"
    _add_cl_body_para(doc, _smartquote(greeting))

    # 12. Blank separator
    _add_cl_blank_separator(doc)

    # 13+. Body paragraphs (with blank separators between them)
    body_paragraphs = cover_letter_content.get("body_paragraphs", [])
    # Include closing as last body paragraph if present
    closing = cover_letter_content.get("closing")
    if closing:
        body_paragraphs = list(body_paragraphs) + [closing]

    for i, para_text in enumerate(body_paragraphs):
        if para_text and para_text.strip():
            _add_cl_body_para(doc, _smartquote(para_text.strip()))
            if i < len(body_paragraphs) - 1:
                _add_cl_blank_separator(doc)

    # N. Blank separator before sign-off
    _add_cl_blank_separator(doc)

    # N+1. "Sincerely," — separate paragraph (default spacing)
    sign_off = cover_letter_content.get("sign_off") or f"Sincerely,\n{name}"
    # Split sign_off into closing phrase and name
    if "\n" in sign_off:
        closing_phrase, sign_name = sign_off.split("\n", 1)
    else:
        closing_phrase = sign_off
        sign_name = name
    closing_phrase = closing_phrase.strip()
    sign_name = sign_name.strip() or name

    sincerely_para = doc.add_paragraph()
    sincerely_para.paragraph_format.space_after = Pt(0)
    sincerely_run = sincerely_para.add_run(_smartquote(closing_phrase))
    sincerely_run.font.name = "Cambria"
    sincerely_run.font.size = Pt(11)

    # N+2. Candidate name — separate paragraph
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(sign_name)
    name_run.font.name = "Cambria"
    name_run.font.size = Pt(11)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Cover letter helpers (PROMPT17 formatting spec)
# ---------------------------------------------------------------------------

_SMARTQUOTE_RE = re.compile(r"(?<=[a-zA-Z])'(?=[a-zA-Z])")


def _smartquote(text: str) -> str:
    """Replace straight apostrophes between letters with \u2019 (right single quote)."""
    return _SMARTQUOTE_RE.sub("\u2019", text)


def _set_cover_letter_defaults(doc: Document) -> None:
    """Set document defaults.

    Cambria 11pt, after=200 twips, line spacing Multiple 1.15.
    """
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Twips(200)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15


def _set_cover_letter_margins(doc: Document) -> None:
    """Set Letter page size, 1-inch margins, and print layout view."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    _set_print_layout_view(doc)


def _add_cover_letter_heading1(doc: Document, name: str) -> None:
    """Add candidate name heading: Calibri 14pt bold #365F91, before=0 after=0."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    ppr = para._p.get_or_add_pPr()
    keep_next = ppr.makeelement(qn("w:keepNext"), {})
    ppr.append(keep_next)
    keep_lines = ppr.makeelement(qn("w:keepLines"), {})
    ppr.append(keep_lines)
    run = para.add_run(name)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x36, 0x5F, 0x91)


def _add_cl_body_para(doc: Document, text: str) -> None:
    """Add a cover letter body paragraph: Cambria 11pt, after=0."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.name = "Cambria"
    run.font.size = Pt(11)


def _add_cl_blank_separator(doc: Document) -> None:
    """Add an empty paragraph that inherits the default spacing (after=200 twips)."""
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Resume helpers
# ---------------------------------------------------------------------------


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15


def _set_margins(doc: Document) -> None:
    """Set Letter page size, 1-inch margins, and print layout view."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    _set_print_layout_view(doc)


def _set_print_layout_view(doc: Document) -> None:
    """Set the document to open in Print Layout view in Word."""
    view_el = doc.settings.element.find(qn("w:view"))
    if view_el is None:
        view_el = doc.settings.element.makeelement(qn("w:view"), {})
        doc.settings.element.append(view_el)
    view_el.set(qn("w:val"), "print")


def _add_section_heading(doc: Document, title: str) -> None:
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(4)
    for run in heading.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.bold = True

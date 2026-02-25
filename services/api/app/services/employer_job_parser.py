"""Employer job document parser.

Extracts structured data from .doc/.docx/.pdf/.txt job descriptions.
"""

import json
import logging
import os
from datetime import datetime
from typing import Any

import anthropic
from docx import Document

logger = logging.getLogger(__name__)

_client: anthropic.Anthropic | None = None


def _get_client() -> anthropic.Anthropic:
    global _client
    if _client is None:
        _client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    return _client


def parse_job_document(file_path: str) -> dict[str, Any]:
    """Parse a job description document and return structured data.

    Args:
        file_path: Path to .doc, .docx, .pdf, or .txt file

    Returns:
        Dictionary with extracted job fields and parsing_confidence.
    """
    from pathlib import Path as _Path

    path = _Path(file_path)
    ext = path.suffix.lower()

    if ext == ".txt":
        full_text = path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".pdf":
        full_text = _extract_pdf_text(file_path)
    elif ext == ".doc":
        full_text = _extract_doc_text(file_path)
    else:
        doc = Document(file_path)
        full_text = _extract_full_text(doc)

    if not full_text.strip():
        return {}

    parsed = _parse_with_claude(full_text)
    return _post_process_job_data(parsed)


def _extract_pdf_text(file_path: str) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())
        return "\n".join(parts)
    except Exception as e:
        logger.error("PDF text extraction failed: %s", e)
        return ""


def _extract_doc_text(file_path: str) -> str:
    """Extract text from a .doc file using Word COM automation (Windows).

    Falls back to LibreOffice conversion if COM is unavailable.
    """
    import os

    abs_path = os.path.abspath(file_path)

    # Try Word COM automation (Windows with MS Word installed)
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(abs_path, ReadOnly=True)
            text = doc.Content.Text
            doc.Close(False)
            return text
        finally:
            word.Quit()
    except Exception as e:
        logger.debug("Word COM extraction failed, trying LibreOffice: %s", e)

    # Fall back to LibreOffice conversion
    try:
        import shutil

        from app.services.doc_converter import convert_doc_to_docx

        from pathlib import Path as _Path

        docx_path = convert_doc_to_docx(_Path(file_path))
        try:
            doc = Document(str(docx_path))
            return _extract_full_text(doc)
        finally:
            shutil.rmtree(docx_path.parent, ignore_errors=True)
    except Exception as e:
        logger.error(".doc extraction failed (no Word or LibreOffice): %s", e)
        return ""


def _extract_full_text(doc: Document) -> str:
    """Extract text from both paragraphs and tables in document order.

    Walks the document body XML to preserve the interleaved ordering of
    paragraphs and tables, so section headers stay next to their tables.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    parts: list[str] = []

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            t_elems = child.findall(".//w:t", ns)
            text = "".join(t.text or "" for t in t_elems).strip()
            if text:
                parts.append(text)
        elif tag == "tbl":
            for tr in child.findall(".//w:tr", ns):
                cells: list[str] = []
                for tc in tr.findall("w:tc", ns):
                    cell_texts = tc.findall(".//w:t", ns)
                    cell_text = "".join(t.text or "" for t in cell_texts).strip()
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    parts.append(" | ".join(cells))

    if not parts:
        # Fallback: read paragraphs only
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

    return "\n".join(parts)


def _parse_with_claude(text: str) -> dict[str, Any]:
    """Use Claude to extract structured job information from text."""
    prompt = (
        "You are a job posting parser. Extract structured "
        "information from this job description.\n\n"
        f"Job Description:\n{text}\n\n"
        "Extract the following fields and return ONLY valid JSON:"
        "\n\n{\n"
        '  "title": "Job title",\n'
        '  "job_id_external": "Job ID, requisition number, '
        "solicitation reference number, posting number, "
        'or vacancy number if mentioned, null otherwise",\n'
        '  "department": '
        '"Department or team if mentioned, null otherwise",\n'
        '  "job_category": "The EXACT category, labor category, '
        "functional area, or job family as written in the "
        "document (e.g. 'Project Management', 'IT', "
        "'Information Technology', 'Administrative'). "
        "Do NOT map to a standard list — use the document's "
        'own wording. null if not mentioned",\n'
        '  "client_company_name": "The client, customer, '
        "end-client, hiring agency, or organization the work "
        "is performed for. For government contracts this is "
        "the agency name (e.g. 'Railroad Commission of Texas',"
        " 'Department of Defense'). "
        'null if not mentioned",\n'
        '  "location": '
        '"Job location if mentioned, null otherwise",\n'
        '  "remote_policy": '
        '"on-site, hybrid, or remote if mentioned, '
        'null otherwise",\n'
        '  "employment_type": "full-time, part-time, contract,'
        ' or internship if mentioned, null otherwise",\n'
        '  "job_type": "permanent, contract, temporary, '
        'or seasonal if mentioned, null otherwise",\n'
        '  "start_date": '
        '"YYYY-MM-DD format if mentioned, null otherwise",\n'
        '  "close_date": '
        '"YYYY-MM-DD format if mentioned, null otherwise",\n'
        '  "description": "COMPLETE job description with '
        'ALL duties and scope. No truncation.",\n'
        '  "requirements": "ALL required qualifications/'
        'skills. Include tables. null if mixed.",\n'
        '  "nice_to_haves": "ALL preferred qualifications. '
        'null if not distinguishable.",\n'
        '  "certifications_required": '
        '["list", "of", "certifications"] or null,\n'
        '  "salary_min": null or integer,\n'
        '  "salary_max": null or integer,\n'
        '  "salary_currency": "USD or other currency code",\n'
        '  "equity_offered": true or false,\n'
        '  "application_email": '
        '"email if mentioned, null otherwise",\n'
        '  "application_url": '
        '"URL to apply if mentioned, null otherwise"\n'
        "}\n\n"
        "Rules:\n"
        "1. Extract ONLY information explicitly stated "
        "in the text\n"
        "2. Do not infer or fabricate information\n"
        "3. Use null for missing fields\n"
        "4. Parse dates carefully — look for "
        '"Application deadline", "Close date", '
        '"Apply by", "Start date", "Expected start", '
        '"Estimated start", "Begin date", '
        '"Submission deadline", etc.\n'
        "5. Return ONLY the JSON object, no other text\n"
        "6. For description and requirements, include ALL "
        "content — do NOT summarize or truncate\n"
        "7. Content from tables is separated by "
        '" | " — include it in the appropriate field\n'
        "8. For job_id_external: use any identifying number "
        "for the position — requisition number, vacancy "
        "number, job ID, solicitation number, or reference "
        "number (e.g., 2025-7642, DIR-CPO-TMP-445). "
        "For government RFRs, use the solicitation number\n"
        "9. For title: use the actual position title "
        "(e.g., 'Senior Business Analyst'), NOT the labor "
        "category or classification level\n"
        "10. For requirements: extract ALL content from "
        "sections labeled Minimum, Required, Mandatory, "
        "or similar — include every row from tables in "
        "these sections as bullet points\n"
        "11. For nice_to_haves: extract ALL content from "
        "sections labeled Preferred, Optional, Desired, "
        "Nice-to-Have, or similar — include every row "
        "from tables in these sections as bullet points\n"
        "12. READ THE ENTIRE DOCUMENT before deciding on "
        "salary. Later sections may override or correct "
        "earlier values. For government contracts: use "
        "the official agency MAX NTE (Not-To-Exceed) rate, "
        "NOT the vendor/contractor proposed rate. "
        'Look for labels like "MAX NTE Rate", '
        '"NTE Rate", "Maximum Rate", "Ceiling Rate". '
        "If the rate is hourly (e.g., $140.00/HR), "
        "store the integer hourly amount in salary_max "
        "(e.g., 140)\n"
        "13. For dates: scan the ENTIRE document including "
        "tables, footers, and appendices. Government RFRs "
        "often list start dates and submission deadlines "
        "in tables or late sections\n"
        "14. For job_category: use the VERBATIM category from "
        "the document. Look for labels like 'Category', "
        "'Labor Category', 'Job Family', 'Functional Area', "
        "'Job Classification'. Do NOT normalize or map to "
        "a standard list — if the document says "
        "'Project Management', return 'Project Management'\n"
        "15. For client_company_name: look for the client, "
        "customer, end-client, or agency name. In government "
        "contracts, this is the agency (e.g. 'Railroad "
        "Commission of Texas', 'Texas DIR'). Look for labels "
        "like 'Client', 'Customer', 'Agency', 'End Client', "
        "'Organization', or the entity issuing the RFR/RFP"
    )

    try:
        client = _get_client()
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}],
        )

        response_text = message.content[0].text.strip()

        # Strip markdown code fences
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]

        return json.loads(response_text.strip())

    except json.JSONDecodeError as e:
        logger.error("Failed to parse Claude JSON response: %s", e)
        return {}
    except anthropic.BadRequestError as e:
        msg = str(e)
        if "credit balance" in msg or "billing" in msg.lower():
            raise RuntimeError("AI service temporarily unavailable (billing). Please try again later.") from e
        logger.error("Claude API call failed: %s", e)
        raise RuntimeError(f"AI parsing failed: {msg}") from e
    except anthropic.APIError as e:
        logger.error("Claude API call failed: %s", e)
        raise RuntimeError("AI service temporarily unavailable. Please try again later.") from e
    except Exception as e:
        logger.error("Claude API call failed: %s", e)
        return {}


def _post_process_job_data(data: dict[str, Any]) -> dict[str, Any]:
    """Clean and validate parsed job data."""
    if not data:
        return data

    # Parse date strings
    for field in ("start_date", "close_date"):
        val = data.get(field)
        if val and isinstance(val, str):
            try:
                data[field] = datetime.strptime(val, "%Y-%m-%d").date()
            except ValueError:
                data[field] = None

    # Ensure certifications is a list
    certs = data.get("certifications_required")
    if certs and not isinstance(certs, list):
        data["certifications_required"] = [certs]

    # Calculate parsing confidence
    required = ["title", "description"]
    optional = ["requirements", "location", "job_category", "department"]
    filled_req = sum(1 for f in required if data.get(f))
    filled_opt = sum(1 for f in optional if data.get(f))

    confidence = (filled_req / len(required)) * 0.7 + (filled_opt / len(optional)) * 0.3
    data["parsing_confidence"] = round(confidence, 2)

    return data

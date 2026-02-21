"""Tests for the form filler service."""

import tempfile
from pathlib import Path

from docx import Document

from app.services.form_filler import (
    _match_skill_to_profile,
    fill_cell_preserving_format,
    fill_form,
)
from app.services.form_parser import parse_form_document


def _make_qualification_docx(path: str) -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Skill / Technology"
    table.rows[0].cells[1].text = "Years of Experience"
    table.rows[1].cells[0].text = "Python"
    table.rows[1].cells[1].text = ""
    table.rows[2].cells[0].text = "AWS"
    table.rows[2].cells[1].text = ""
    doc.save(path)


def test_fill_skills_matrix():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        template = f.name

    _make_qualification_docx(template)
    parsed = parse_form_document(template, job_id=1)

    output = template.replace(".docx", "_filled.docx")
    profile = {
        "basics": {"name": "Test User"},
        "skills": ["Python", "AWS"],
        "skill_years": {
            "Python": {"years_experience": 5, "years_experience_source": "parsed"},
            "AWS": {"years_experience": 3, "years_experience_source": "manual"},
        },
        "references": [],
    }

    result = fill_form(template, parsed, profile, output)

    assert result["filled_count"] >= 1
    assert Path(output).exists()

    # Verify the filled DOCX
    filled_doc = Document(output)
    table = filled_doc.tables[0]
    python_years = table.rows[1].cells[1].text
    assert python_years == "5"
    aws_years = table.rows[2].cells[1].text
    assert aws_years == "3"

    Path(template).unlink(missing_ok=True)
    Path(output).unlink(missing_ok=True)


def test_skill_matching_exact():
    skill_years = {
        "Python": {"years_experience": 5, "years_experience_source": "parsed"},
    }
    result = _match_skill_to_profile("Python", skill_years, ["python"])
    assert result is not None
    assert result["years_experience"] == 5


def test_skill_matching_substring():
    skill_years = {
        "Python": {"years_experience": 5, "years_experience_source": "parsed"},
    }
    result = _match_skill_to_profile("Python programming", skill_years, [])
    assert result is not None
    assert result["years_experience"] == 5


def test_skill_matching_no_match():
    result = _match_skill_to_profile("Haskell", {}, ["python", "java"])
    assert result is None


def test_fill_cell_preserving_format():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].add_run("placeholder")

    fill_cell_preserving_format(cell, "new value")

    assert cell.paragraphs[0].runs[0].text == "new value"

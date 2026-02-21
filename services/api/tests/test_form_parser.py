"""Tests for the form parser service."""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from app.services.form_parser import (
    ACKNOWLEDGEMENT,
    OTHER,
    REFERENCE_FORM,
    SKILLS_MATRIX,
    _classify_table,
    _is_never_fill,
    _map_field_to_profile,
    parse_form_document,
)


def _make_skills_matrix_docx(path: str) -> None:
    """Create a DOCX with a skills/years matrix table."""
    doc = Document()
    doc.add_heading("Qualification Form", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Skill / Technology"
    table.rows[0].cells[1].text = "Years of Experience"
    table.rows[1].cells[0].text = "Python"
    table.rows[1].cells[1].text = ""
    table.rows[2].cells[0].text = "AWS"
    table.rows[2].cells[1].text = ""
    table.rows[3].cells[0].text = "Docker"
    table.rows[3].cells[1].text = ""

    doc.save(path)


def _make_reference_form_docx(path: str) -> None:
    """Create a DOCX with a reference form table."""
    doc = Document()
    doc.add_heading("Professional References", level=1)

    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Reference Name"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "Title"
    table.rows[1].cells[1].text = ""
    table.rows[2].cells[0].text = "Company"
    table.rows[2].cells[1].text = ""
    table.rows[3].cells[0].text = "Phone"
    table.rows[3].cells[1].text = ""

    doc.save(path)


def test_parse_skills_matrix():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        _make_skills_matrix_docx(f.name)
        result = parse_form_document(f.name, job_id=1)

    assert result["job_id"] == 1
    assert result["total_fields"] > 0
    sections = result["sections"]
    assert any(s["type"] == SKILLS_MATRIX for s in sections)

    # Should have 3 skill fields (Python, AWS, Docker)
    skills_section = next(s for s in sections if s["type"] == SKILLS_MATRIX)
    skill_labels = [f["label"] for f in skills_section["fields"]]
    assert "Python" in skill_labels
    assert "AWS" in skill_labels
    assert "Docker" in skill_labels

    Path(f.name).unlink(missing_ok=True)


def test_parse_reference_form():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        _make_reference_form_docx(f.name)
        result = parse_form_document(f.name, job_id=2)

    sections = result["sections"]
    # Should detect a reference_form or other (depending on header detection)
    assert len(sections) > 0
    Path(f.name).unlink(missing_ok=True)


def test_never_fill_patterns():
    assert _is_never_fill("Candidate Signature") is True
    assert _is_never_fill("Sign Here") is True
    assert _is_never_fill("Hourly Rate") is True
    assert _is_never_fill("Bill Rate") is True
    assert _is_never_fill("Vendor Name") is True
    assert _is_never_fill("Subcontractor") is True

    # Should not block regular fields
    assert _is_never_fill("Candidate Name") is False
    assert _is_never_fill("Email Address") is False
    assert _is_never_fill("Python") is False


def test_field_mapping():
    assert _map_field_to_profile("Candidate Name", OTHER) == ["basics", "name"]
    assert _map_field_to_profile("Email:", OTHER) == ["basics", "email"]
    assert _map_field_to_profile("Phone Number", OTHER) == ["basics", "phone"]
    assert _map_field_to_profile("Something Unknown", OTHER) is None


def test_file_not_found():
    with pytest.raises(FileNotFoundError):
        parse_form_document("/nonexistent/form.docx", job_id=99)

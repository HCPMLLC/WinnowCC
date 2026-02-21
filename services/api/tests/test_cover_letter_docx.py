"""Tests for cover letter DOCX formatting compliance (PROMPT17 spec)."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips

from app.services.docx_builder import (
    _smartquote,
    build_cover_letter_docx,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

CANDIDATE_BASICS = {
    "name": "Jane Doe",
    "email": "jane@example.com",
    "phone": "(555) 123-4567",
    "location": "Denver, CO",
    "linkedin": "linkedin.com/in/janedoe",
}

COVER_LETTER_CONTENT = {
    "greeting": "Dear Hiring Manager,",
    "body_paragraphs": [
        "I'm excited to apply for the role. My background aligns well.",
        "At Acme Corp I managed cross-functional teams and delivered results.",
    ],
    "closing": "I look forward to discussing how I can contribute to your team.",
    "sign_off": "Sincerely,\nJane Doe",
}

JOB_INFO = {
    "company": "TechCorp",
    "title": "Project Manager",
    "hiring_manager": "John Smith",
}


@pytest.fixture()
def cover_letter_path(tmp_path: Path) -> Path:
    out = tmp_path / "cover.docx"
    build_cover_letter_docx(
        cover_letter_content=COVER_LETTER_CONTENT,
        candidate_basics=CANDIDATE_BASICS,
        job_info=JOB_INFO,
        output_path=str(out),
    )
    return out


@pytest.fixture()
def cover_letter_doc(cover_letter_path: Path) -> Document:
    return Document(str(cover_letter_path))


# ---------------------------------------------------------------------------
# 1. Page margins
# ---------------------------------------------------------------------------


def test_page_margins(cover_letter_doc: Document) -> None:
    section = cover_letter_doc.sections[0]
    # 1 inch = 914400 EMU = 1440 twips (DXA)
    assert section.top_margin == Inches(1)
    assert section.bottom_margin == Inches(1)
    # 1 inch (matches _set_cover_letter_margins)
    assert section.left_margin == Inches(1)
    assert section.right_margin == Inches(1)


# ---------------------------------------------------------------------------
# 2. Heading font — Calibri 14pt bold #365F91
# ---------------------------------------------------------------------------


def test_heading_font(cover_letter_doc: Document) -> None:
    first_para = cover_letter_doc.paragraphs[0]
    assert first_para.text == "Jane Doe"
    run = first_para.runs[0]
    assert run.bold is True
    assert run.font.name == "Calibri"
    assert run.font.size == Pt(14)
    assert run.font.color.rgb == RGBColor(0x36, 0x5F, 0x91)


# ---------------------------------------------------------------------------
# 3. Body font — Cambria 11pt
# ---------------------------------------------------------------------------


def test_body_font(cover_letter_doc: Document) -> None:
    # Contact line is paragraph index 1
    contact_para = cover_letter_doc.paragraphs[1]
    assert "\u2022" in contact_para.text  # bullet separator
    run = contact_para.runs[0]
    assert run.font.name == "Cambria"
    assert run.font.size == Pt(11)


def test_document_default_font(cover_letter_doc: Document) -> None:
    style = cover_letter_doc.styles["Normal"]
    assert style.font.name == "Cambria"
    assert style.font.size == Pt(11)


# ---------------------------------------------------------------------------
# 4. Paragraph structure — correct count and blank separators
# ---------------------------------------------------------------------------


def test_paragraph_structure(cover_letter_doc: Document) -> None:
    paras = cover_letter_doc.paragraphs
    texts = [p.text for p in paras]

    # Expected structure:
    # 0: "Jane Doe" (heading)
    # 1: contact line with bullets
    # 2: blank separator
    # 3: date
    # 4: blank separator
    # 5: "John Smith" (hiring manager)
    # 6: "TechCorp"
    # 7: blank separator
    # 8: "Dear Hiring Manager,"
    # 9: blank separator
    # 10: body para 1
    # 11: blank separator
    # 12: body para 2
    # 13: blank separator
    # 14: closing para
    # 15: blank separator
    # 16: "Sincerely,"
    # 17: "Jane Doe"

    assert texts[0] == "Jane Doe"
    assert texts[2] == ""  # blank separator after contact
    assert texts[5] == "John Smith"
    assert texts[6] == "TechCorp"
    assert texts[16] == "Sincerely,"
    assert texts[17] == "Jane Doe"


# ---------------------------------------------------------------------------
# 5. Smart quotes
# ---------------------------------------------------------------------------


def test_smartquote_function() -> None:
    assert _smartquote("I'm excited") == "I\u2019m excited"
    assert _smartquote("can't won't") == "can\u2019t won\u2019t"
    # Should not affect non-apostrophe quotes
    assert _smartquote("'hello'") == "'hello'"


def test_smart_quotes_in_body(cover_letter_doc: Document) -> None:
    # The first body paragraph contains "I'm" which should be smart-quoted
    body_texts = [p.text for p in cover_letter_doc.paragraphs if "excited" in p.text]
    assert len(body_texts) == 1
    assert "\u2019" in body_texts[0]  # curly apostrophe


# ---------------------------------------------------------------------------
# 6. Contact line separator — bullet not pipe
# ---------------------------------------------------------------------------


def test_contact_line_uses_bullet(cover_letter_doc: Document) -> None:
    contact_para = cover_letter_doc.paragraphs[1]
    assert "\u2022" in contact_para.text
    assert "|" not in contact_para.text


# ---------------------------------------------------------------------------
# 7. Sign-off — separate paragraphs
# ---------------------------------------------------------------------------


def test_signoff_separate_paragraphs(cover_letter_doc: Document) -> None:
    paras = cover_letter_doc.paragraphs
    texts = [p.text for p in paras]
    sincerely_idx = texts.index("Sincerely,")
    assert texts[sincerely_idx + 1] == "Jane Doe"


# ---------------------------------------------------------------------------
# 8. Template path produces same formatting
# ---------------------------------------------------------------------------


def test_template_path_formatting(tmp_path: Path) -> None:
    """Verify _build_cover_letter_doc delegates to build_cover_letter_docx."""
    from app.services.tailor import _build_cover_letter_doc

    # Create a mock Job object
    job = MagicMock()
    job.title = "Software Engineer"
    job.company = "TestCorp"
    job.description_text = "Looking for a software engineer with Python experience."
    job.hiring_manager_name = None

    profile_json = {
        "basics": {
            "name": "Test User",
            "email": "test@example.com",
            "phone": "(555) 000-0000",
            "location": "Austin, TX",
        },
        "skills": ["Python", "JavaScript"],
        "experience": [
            {
                "title": "Developer",
                "company": "OldCo",
                "duties": ["Built stuff"],
            }
        ],
    }

    out = tmp_path / "template_cover.docx"

    with patch(
        "app.services.cover_letter_generator.generate_cover_letter_text",
        return_value=(
            "Dear Hiring Manager,\n\n"
            "I am excited to apply for the Software Engineer role at TestCorp.\n\n"
            "My background in Python makes me a strong fit.\n\n"
            "I look forward to contributing to your team.\n\n"
            "Sincerely,\nTest User"
        ),
    ):
        _build_cover_letter_doc(out, job, profile_json)

    doc = Document(str(out))

    # Verify it uses PROMPT17 formatting
    section = doc.sections[0]
    assert section.top_margin == Inches(1)
    assert section.left_margin == Inches(1)

    # Heading: first para should be candidate name
    assert doc.paragraphs[0].text == "Test User"
    heading_run = doc.paragraphs[0].runs[0]
    assert heading_run.font.name == "Calibri"
    assert heading_run.font.size == Pt(14)
    assert heading_run.bold is True

    # Contact line uses bullet
    contact = doc.paragraphs[1].text
    assert "\u2022" in contact


# ---------------------------------------------------------------------------
# 9. Default line spacing
# ---------------------------------------------------------------------------


def test_default_line_spacing(cover_letter_doc: Document) -> None:
    style = cover_letter_doc.styles["Normal"]
    # Multiple 1.15 (matches _set_cover_letter_defaults)
    assert style.paragraph_format.line_spacing == 1.15


def test_default_space_after(cover_letter_doc: Document) -> None:
    style = cover_letter_doc.styles["Normal"]
    assert style.paragraph_format.space_after == Twips(200)


# ---------------------------------------------------------------------------
# 10. Content paragraph after=0
# ---------------------------------------------------------------------------


def test_body_para_spacing(cover_letter_doc: Document) -> None:
    """Content paragraphs (non-blank) should have after=0."""
    contact_para = cover_letter_doc.paragraphs[1]  # contact line
    assert contact_para.paragraph_format.space_after == Pt(0)


# ---------------------------------------------------------------------------
# 11. Address block with hiring manager
# ---------------------------------------------------------------------------


def test_address_block_with_address(tmp_path: Path) -> None:
    out = tmp_path / "cover_addr.docx"
    job_info = {
        "company": "BigCo",
        "title": "PM",
        "hiring_manager": "Alice Lee",
        "address": "123 Main St\nSuite 100\nNew York, NY 10001",
    }
    build_cover_letter_docx(
        cover_letter_content=COVER_LETTER_CONTENT,
        candidate_basics=CANDIDATE_BASICS,
        job_info=job_info,
        output_path=str(out),
    )
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Alice Lee" in texts
    assert "BigCo" in texts
    assert "123 Main St" in texts
    assert "Suite 100" in texts
    assert "New York, NY 10001" in texts
